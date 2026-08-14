import asyncio
"""
小赛助手 v2 - FastAPI 后端服务
赛乐赛瘦身产品销售助手
"""

import os
import sys
import json
import webbrowser
from pathlib import Path
from datetime import datetime
from contextlib import asynccontextmanager

# 将 backend 目录加入路径
BACKEND_DIR = Path(__file__).resolve().parent
BASE_DIR = BACKEND_DIR.parent
sys.path.insert(0, str(BACKEND_DIR))

from fastapi import FastAPI, HTTPException, Query, Body, Form, UploadFile, File, Request
from typing import List
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import HTMLResponse, JSONResponse, FileResponse, StreamingResponse, Response
import uvicorn

from database import (init_db, list_customers, get_customer, create_customer,
                      update_customer, delete_customer, group_customers,
                      get_messages, add_message, delete_message, get_db,
                      export_all, import_data, parse_customer_text,
                      scan_folder_for_index, image_index_sync, image_index_scan_incremental, get_image_index_status,
                      get_pending_images, update_image_result, search_images,
                      search_images_hybrid, rebuild_image_vectors, rebuild_category_profile_vectors,
                      _image_reindex_progress_callback,
                      image_index_clear, get_image_detail, get_case_group, match_images_for_script,
                      auto_group_cases, assign_case_group, get_ungrouped_images,
                      get_category_images, delete_image_category,
                      get_all_subcategories, update_subcategory_profile,
                      delete_subcategory, get_subcategory_images,
                      update_image_fields, soft_delete_image,
                      batch_delete_images, batch_tag_images,
                      add_image_record, copy_image_to_category,
                      increment_image_use_count)
from message_parser import parse_and_split_messages, group_consecutive_messages
from effective_scripts import (add_effective_script, list_effective_scripts,
                               get_effective_script_stats, delete_effective_script,
                               search_effective_scripts_by_scenario, dedup_effective_scripts,
                               update_script_vector_status, revectorize_script,
                               update_effective_script)
from faq import (list_faqs, get_faq, add_faq, update_faq, delete_faq)
from knowledge import (list_documents, get_document, add_document,
                       delete_document, search_knowledge, reindex_all,
                       get_knowledge_status, get_document_content,
                       is_embedding_available,
                       extract_text_from_file, save_uploaded_file,
                       add_customer_image, get_customer_images,
                       search_customer_images, delete_customer_image)
from llm_client import generate_script, generate_script_stream
import config_manager as cfg

# ===== 生成图片静态服务 =====
GENERATED_IMAGES_DIR = (cfg.DATA_DIR / "generated_images").resolve()

# 占位 SVG，当图片文件不存在时返回（避免前端 404 报错）
PLACEHOLDER_SVG = '<svg xmlns="http://www.w3.org/2000/svg" width="200" height="200" viewBox="0 0 200 200"><rect width="200" height="200" fill="#f0f0f0"/><text x="100" y="95" text-anchor="middle" fill="#ccc" font-size="48" font-family="sans-serif">?</text><text x="100" y="125" text-anchor="middle" fill="#aaa" font-size="12" font-family="sans-serif">\u56fe\u7247\u4e0d\u5b58\u5728</text></svg>'


@asynccontextmanager
async def lifespan(app: FastAPI):
    """应用生命周期"""
    # 启动时初始化
    init_db()
    print(f"[小赛助手 v2] 启动成功")
    print(f"   端口: {cfg.load_config()['app']['port']}")
    print(f"   数据目录: {cfg.DATA_DIR}")
    
    # ===== 启动时 API 健康检查 =====
    try:
        from llm_client import get_llm_client
        api_config = get_llm_client()
        api_key = api_config.get("api_key", "")
        base_url = api_config.get("base_url", "")
        model = api_config.get("model", "")
        
        if not api_key or api_key == "你的API_KEY":
            print("! 警告: API KEY 未配置，请在设置页面填入 API KEY")
            print(f"   API地址: {base_url}")
            print(f"   模型: {model}")
        else:
            key_preview = api_key[:8] + "****" + api_key[-4:]
            print(f"   API密钥: {key_preview}")
            print(f"   API地址: {base_url}")
            print(f"   模型: {model}")
            # 测试 API 连通性
            try:
                import httpx
                test_url = f"{base_url.rstrip('/')}/chat/completions"
                test_headers = {
                    "Authorization": f"Bearer {api_key}",
                    "Content-Type": "application/json"
                }
                test_payload = {
                    "model": model,
                    "messages": [{"role": "user", "content": "Hi"}],
                    "max_tokens": 5
                }
                with httpx.Client(timeout=15) as client:
                    test_resp = client.post(test_url, json=test_payload, headers=test_headers)
                    if test_resp.status_code == 200:
                        print(f"   ✅ API 连通性检查通过")
                    else:
                        error_body = test_resp.text[:100]
                        print(f"   ⚠️ API 返回异常 (HTTP {test_resp.status_code}): {error_body}")
                        print(f"   请检查 API 地址和密钥是否正确")
            except httpx.ConnectError:
                print(f"   ❌ API 连接失败: 无法连接到 {base_url}")
                print(f"   请检查网络连接和 API 地址配置")
            except httpx.TimeoutException:
                print(f"   ❌ API 连接超时: {base_url} 响应超过15秒")
                print(f"   请检查网络连接或更换 API 地址")
            except Exception as e:
                print(f"   ❌ API 检查异常: {str(e)[:100]}")
    except Exception as e:
        print(f"! API 配置检查失败: {str(e)[:100]}")
    
    yield
    # 关闭时清理
    print("服务已停止")


# ===== 图片索引进度 =====
_scan_progress = {
    "total": 0,
    "pending": 0,
    "done": 0,
    "failed": 0,
    "skipped": 0,
    "current": None,
    "status": "idle",
    "started_at": None,
    "completed_at": None,
    "cancel_flag": False,
    "pause_flag": False,
    "task_type": None,
    "base_path": "",
    "categories": [],  # [{name, total, done, failed, status}]
    "current_category": "",
    "total_categories": 0,
    "done_categories": 0,
}
_model_stats = {}  # {model_id: {"success":0, "fail":0, "consecutive_fail":0, "last_try":0}}
_scan_lock = None  # asyncio.Lock 会在异步上下文中创建
_scan_lock_obj = None

async def get_scan_lock():
    global _scan_lock_obj
    if _scan_lock_obj is None:
        import asyncio
        _scan_lock_obj = asyncio.Lock()
    return _scan_lock_obj

app = FastAPI(title="小赛助手 v2", lifespan=lifespan)

# CORS 允许跨域
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ==================== API 端点 ====================

@app.get("/api/health")
async def health():
    """健康检查"""
    return {"status": "ok", "time": datetime.now().strftime("%Y-%m-%d %H:%M:%S")}


# ===== 配置管理 =====

@app.get("/api/config")
async def get_config():
    """获取LLM配置（隐藏API密钥中间部分）"""
    config = cfg.get_llm_config()
    # 脱敏显示
    masked = dict(config)
    key = masked.get("api_key", "")
    if len(key) > 8:
        masked["api_key"] = key[:8] + "****" + key[-4:]
    return masked


@app.get("/api/app-info")
async def get_app_info():
    """获取应用信息（路径等）"""
    return {
        "base_dir": str(BASE_DIR),
        "myalbums_dir": str(BASE_DIR / "MyAlbums"),
        "data_dir": str(cfg.DATA_DIR)
    }


@app.put("/api/config")
async def update_config(data: dict = Body(...)):
    """更新LLM配置"""
    provider = data.get("provider", "")
    api_key = data.get("api_key", "")
    model = data.get("model", "")
    base_url = data.get("base_url", "")

    if not api_key:
        raise HTTPException(status_code=400, detail="API Key 不能为空")
    if not base_url:
        raise HTTPException(status_code=400, detail="Base URL 不能为空")

    updated = cfg.update_llm_config(provider, api_key, model, base_url)
    # 脱敏
    masked = dict(updated)
    key = masked.get("api_key", "")
    if len(key) > 8:
        masked["api_key"] = key[:8] + "****" + key[-4:]
    return {"status": "ok", "config": masked}


# ===== API过期提醒 =====

@app.get("/api/check-expiry")
async def api_check_expiry():
    """检查所有已配置模型的API密钥是否有效"""
    import httpx

    # 主模型
    from llm_client import get_llm_client
    main_config = get_llm_client()
    api_key = main_config.get("api_key", "")
    base_url = main_config.get("base_url", "")
    model = main_config.get("model", "")

    result = {
        "configured": bool(api_key and api_key != "你的API_KEY"),
        "base_url": base_url,
        "model": model,
        "reachable": False,
        "status": "unconfigured",
        "message": "",
        "models": []
    }

    if not result["configured"]:
        result["message"] = "API密钥未配置"
    else:
        try:
            test_url = f"{base_url.rstrip(chr(47))}/chat/completions"
            test_headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
            test_payload = {"model": model, "messages": [{"role": "user", "content": "Hi"}], "max_tokens": 5}
            with httpx.Client(timeout=15) as client:
                test_resp = client.post(test_url, json=test_payload, headers=test_headers)
                if test_resp.status_code == 200:
                    result["reachable"] = True
                    result["status"] = "ok"
                    result["message"] = "API连接正常"
                elif test_resp.status_code in (401, 403):
                    result["reachable"] = False
                    result["status"] = "error"
                    result["message"] = f"API返回异常 (HTTP {test_resp.status_code})"
                else:
                    result["reachable"] = False
                    result["status"] = "warning"
                    result["message"] = f"API返回状态码 {test_resp.status_code}"
        except httpx.ConnectError:
            result["reachable"] = False
            result["status"] = "error"
            result["message"] = "无法连接到API服务器"
        except httpx.TimeoutException:
            result["reachable"] = False
            result["status"] = "error"
            result["message"] = "API连接超时"
        except Exception as e:
            result["reachable"] = False
            result["status"] = "error"
            result["message"] = f"检测失败: {str(e)[:80]}"

    # 检测所有已配置模型
    from config_manager import load_config
    full_config = load_config()
    all_model_dict = full_config.get("models", {}).get("list", {})
    result["models"] = []
    for mid, m_data in all_model_dict.items():
        murl = m_data.get("base_url", "")
        mmodel = m_data.get("model", "")
        mkey = m_data.get("api_key", "")
        mprovider = m_data.get("provider", "")
        mcats = m_data.get("categories", [])
        m_api_mode = m_data.get("api_mode", "openai")
        status = "unknown"
        msg = ""
        if murl and mmodel and mkey and mkey != "你的API_KEY":
            try:
                is_emb = "embedding" in mcats or m_api_mode in ("dashscope_text", "dashscope_multimodal")
                if is_emb:
                    if m_api_mode == "dashscope_multimodal":
                        test_payload2 = {"model": mmodel, "input": {"contents": [{"text": "测试"}]}, "parameters": {"dimension": 256}}
                    elif m_api_mode == "dashscope_text":
                        test_payload2 = {"model": mmodel, "input": {"texts": ["测试"]}}
                    else:
                        test_payload2 = {"input": "测试", "model": mmodel, "dimensions": 1024}
                        # OpenAI 模式自动补 /embeddings
                        if not murl.rstrip("/").endswith("/embeddings"):
                            murl = murl.rstrip("/") + "/embeddings"
                    test_url2 = murl
                else:
                    test_url2 = f"{murl.rstrip(chr(47))}/chat/completions"
                    test_payload2 = {"model": mmodel, "messages": [{"role": "user", "content": "Hi"}], "max_tokens": 5}
                test_headers2 = {"Authorization": f"Bearer {mkey}", "Content-Type": "application/json"}
                with httpx.Client(timeout=8) as client:
                    r = client.post(test_url2, json=test_payload2, headers=test_headers2)
                    if r.status_code == 200:
                        status = "ok"
                        msg = "正常"
                    elif r.status_code in (401, 403):
                        status = "error"
                        msg = f"密钥无效 (HTTP {r.status_code})"
                    else:
                        status = "warning"
                        msg = f"状态码 {r.status_code}"
            except Exception:
                status = "error"
                msg = "不可达"
        result["models"].append({
            "id": mid,
            "model": mmodel,
            "provider": mprovider,
            "categories": mcats,
            "base_url": murl[:40],
            "status": status,
            "message": msg
        })

    return result
# ===== 客户管理 =====

@app.get("/api/customers")
async def api_list_customers(search: str = Query("", description="搜索关键字")):
    """获取客户列表"""
    customers = list_customers(search)
    return {"customers": customers}


@app.post("/api/customers")
async def api_create_customer(data: dict = Body(...)):
    """创建客户"""
    required = ["name"]
    for field in required:
        if not data.get(field):
            raise HTTPException(status_code=400, detail=f"{field} 不能为空")

    customer = create_customer(
        name=data.get("name", ""),
        title=data.get("title", ""),
        age=data.get("age", ""),
        height=data.get("height", ""),
        weight=data.get("weight", ""),
        target_weight=data.get("target_weight", ""),
        purchase=data.get("purchase", ""),
        purchase_history=data.get("purchase_history", "[]"),
        customer_type=data.get("customer_type", ""),
        remark=data.get("remark", "")
    )
    return {"customer": customer}


@app.post("/api/customers/batch")
async def api_batch_create_customers(data: dict = Body(...)):
    """批量导入客户（支持文本格式和JSON数组）"""
    results = {"success": 0, "failed": 0, "errors": []}
    
    # 文本格式：每行一个客户，格式如 "姓名身高.体重.年龄" 或 "姓名 备注"
    text = data.get("text", "")
    if text:
        lines = [l.strip() for l in text.split("\n") if l.strip()]
        for line in lines:
            try:
                parsed = parse_customer_text(line)
                if not parsed.get("name"):
                    results["failed"] += 1
                    results["errors"].append(f"无法解析: {line[:30]}")
                    continue
                customer = create_customer(
                    name=parsed.get("name", ""),
                    age=parsed.get("age", ""),
                    height=parsed.get("height", ""),
                    weight=parsed.get("weight", ""),
                    remark=parsed.get("remark", "")
                )
                results["success"] += 1
            except Exception as e:
                results["failed"] += 1
                results["errors"].append(f"{line[:20]}: {str(e)[:50]}")
    
    # JSON数组格式
    customers = data.get("customers", [])
    if customers:
        for c in customers:
            try:
                if not c.get("name"):
                    results["failed"] += 1
                    results["errors"].append(f"缺少姓名: {str(c)[:50]}")
                    continue
                customer = create_customer(
                    name=c.get("name", ""),
                    title=c.get("title", ""),
                    age=c.get("age", ""),
                    height=c.get("height", ""),
                    weight=c.get("weight", ""),
                    target_weight=c.get("target_weight", ""),
                    purchase=c.get("purchase", ""),
                    purchase_history=c.get("purchase_history", "[]"),
                    customer_type=c.get("customer_type", ""),
                    remark=c.get("remark", "")
                )
                results["success"] += 1
            except Exception as e:
                results["failed"] += 1
                results["errors"].append(f"{c.get('name', '?')}: {str(e)[:50]}")
    
    return {"status": "ok", "results": results}


@app.get("/api/customers/grouped")
async def api_grouped_customers():
    """获取分组客户列表"""
    from database import group_customers
    groups = group_customers()
    return {"groups": groups}


@app.get("/api/follow-up")
async def api_follow_up(days: int = Query(3, description="沉默天数阈值")):
    """获取需要跟进的沉默客户列表（忽略已标记忽略的）"""
    from datetime import datetime, timedelta
    from database import get_db
    from config_manager import load_config, save_config
    
    # 读取已忽略列表
    config = load_config()
    ignored_ids = set(config.get("follow_up_ignored", []))
    
    conn = get_db()
    try:
        threshold = (datetime.now() - timedelta(days=days)).strftime("%Y-%m-%d %H:%M:%S")
        from database import get_messages_db
        msg_conn = get_messages_db()
        cursor = conn.execute("""
            SELECT c.id, c.name, c.title, c.purchase
            FROM customers c
        """)
        all_customers = [dict(r) for r in cursor.fetchall()]
        # 单独查询消息时间戳（消息在单独的消息库）
        for c in all_customers:
            mc = msg_conn.execute("SELECT MAX(timestamp) FROM messages WHERE customer_id=?", (c["id"],))
            last_ts = mc.fetchone()[0]
            c["last_msg_time"] = last_ts
            if last_ts:
                mc2 = msg_conn.execute("SELECT content FROM messages WHERE customer_id=? ORDER BY timestamp DESC LIMIT 1", (c["id"],))
                last_content = mc2.fetchone()
                c["last_msg_preview"] = (last_content[0] or "")[:60] if last_content else ""

        # 检查今天有没有消息
        today_start = datetime.now().strftime("%Y-%m-%d 00:00:00")
        
        silent = []
        for c in all_customers:
            if c.get("id") in ignored_ids:
                continue
            if not c.get("last_msg_time") or c["last_msg_time"] < threshold:
                s_days = (datetime.now() - datetime.strptime(c["last_msg_time"], "%Y-%m-%d %H:%M:%S")).days if c.get("last_msg_time") else 999
                silent.append({
                    "id": c["id"],
                    "name": c["name"],
                    "title": c.get("title", ""),
                    "purchase": c.get("purchase", ""),
                    "last_msg_time": c.get("last_msg_time", "无记录"),
                    "silent_days": s_days,
                    "last_msg_preview": (c.get("last_msg_preview") or "")[:60]
                })
        # 按沉默天数降序排列（最久的在最上面）
        silent.sort(key=lambda x: x["silent_days"], reverse=True)
        return {"count": len(silent), "customers": silent, "threshold_days": days}
    finally:
        conn.close()


@app.post("/api/follow-up/ignore")
async def api_follow_up_ignore(data: dict = Body(...)):
    """忽略某个客户的跟进提醒"""
    customer_id = data.get("customer_id")
    if not customer_id:
        raise HTTPException(status_code=400, detail="缺少 customer_id")
    from config_manager import load_config, save_config
    config = load_config()
    ignored = set(config.get("follow_up_ignored", []))
    ignored.add(customer_id)
    config["follow_up_ignored"] = list(ignored)
    save_config(config)
    return {"status": "ok"}


@app.post("/api/follow-up/undo-ignore")
async def api_follow_up_undo_ignore(data: dict = Body(...)):
    """撤销忽略某个客户的跟进提醒"""
    customer_id = data.get("customer_id")
    if not customer_id:
        raise HTTPException(status_code=400, detail="缺少 customer_id")
    from config_manager import load_config, save_config
    config = load_config()
    ignored = set(config.get("follow_up_ignored", []))
    ignored.discard(customer_id)
    config["follow_up_ignored"] = list(ignored)
    save_config(config)
    return {"status": "ok"}



@app.get("/api/customers/{customer_id}/profile")
async def api_customer_profile(customer_id: int):
    """获取客户画像"""
    customer = get_customer(customer_id)
    if not customer:
        raise HTTPException(status_code=404, detail="客户不存在")

    from database import get_db
    conn = get_db()
    try:
        # 消息统计
        cursor = conn.execute("SELECT COUNT(*) as cnt FROM messages WHERE customer_id=?", (customer_id,))
        msg_count = cursor.fetchone()["cnt"]

        cursor = conn.execute("SELECT MIN(timestamp) as first, MAX(timestamp) as last FROM messages WHERE customer_id=?", (customer_id,))
        times = cursor.fetchone()
        first_msg = times["first"] or "无"
        last_msg = times["last"] or "无"

        # 统计角色分布
        cursor = conn.execute("SELECT role, COUNT(*) as cnt FROM messages WHERE customer_id=? GROUP BY role", (customer_id,))
        role_stats = {r["role"]: r["cnt"] for r in cursor.fetchall()}

        # 消息总数
        user_msgs = role_stats.get("user", 0)
        assistant_msgs = role_stats.get("assistant", 0)

        # 最近10条消息摘要
        cursor = conn.execute("SELECT content, role, timestamp FROM messages WHERE customer_id=? ORDER BY timestamp DESC LIMIT 10", (customer_id,))
        recent = [dict(r) for r in cursor.fetchall()]

        # 图片数
        cursor = conn.execute("SELECT COUNT(*) as cnt FROM customer_images WHERE customer_id=?", (customer_id,))
        img_count = cursor.fetchone()["cnt"] if cursor else 0

        # 获取画像字段
        cursor = conn.execute("""
            SELECT emotion_score, churn_risk, lifecycle_stage, rfm_tier, 
                   rfm_r_days, rfm_f_count, rfm_m_amount, value_score, ai_analysis_at
            FROM customers WHERE id=?
        """, (customer_id,))
        profile_row = cursor.fetchone()

        profile = {
            "customer_id": customer_id,
            "name": customer.get("name", ""),
            "title": customer.get("title", ""),
            "age": customer.get("age", ""),
            "height": customer.get("height", ""),
            "weight": customer.get("weight", ""),
            "purchase": customer.get("purchase", ""),
            "remark": customer.get("remark", ""),
            "msg_count": msg_count,
            "user_msg_count": user_msgs,
            "assistant_msg_count": assistant_msgs,
            "first_msg_time": first_msg,
            "last_msg_time": last_msg,
            "image_count": img_count,
            "recent_messages": recent[:5],
            # 画像字段
            "emotion_score": profile_row[0] if profile_row else 0,
            "churn_risk": profile_row[1] if profile_row else 'low',
            "lifecycle_stage": profile_row[2] if profile_row else 'new',
            "rfm_tier": profile_row[3] if profile_row else 'D',
            "rfm_r_days": profile_row[4] if profile_row else None,
            "rfm_f_count": profile_row[5] if profile_row else 0,
            "rfm_m_amount": profile_row[6] if profile_row else 0,
            "value_score": profile_row[7] if profile_row else 0,
            "ai_analysis_at": profile_row[8] if profile_row else None,
        }
        return profile
    finally:
        conn.close()


@app.post("/api/customers/{customer_id}/analyze")
async def api_analyze_customer(customer_id: int):
    """分析客户画像并更新数据库"""
    from customer_profile import analyze_customer_profile
    profile = analyze_customer_profile(customer_id)
    if not profile:
        raise HTTPException(status_code=404, detail="客户不存在")
    return {"status": "ok", "profile": profile}


@app.get("/api/customers/{customer_id}/profile/analysis")
async def api_get_customer_analysis(customer_id: int):
    """获取客户画像分析结果（不更新）"""
    from customer_profile import analyze_customer_profile
    from database import get_customer
    customer = get_customer(customer_id)
    if not customer:
        raise HTTPException(status_code=404, detail="客户不存在")
    profile = analyze_customer_profile(customer_id)
    return {"profile": profile}


@app.post("/api/customers/batch-analyze")
async def api_batch_analyze_customers(limit: int = 50):
    """批量分析客户画像"""
    from customer_profile import batch_analyze_customers
    result = batch_analyze_customers(limit)
    return {"status": "ok", "result": result}
# ===== 客户图片 =====

@app.get("/api/customers/{customer_id}")
async def api_get_customer(customer_id: int):
    """获取单个客户"""
    customer = get_customer(customer_id)
    if not customer:
        raise HTTPException(status_code=404, detail="客户不存在")
    return {"customer": customer}


@app.put("/api/customers/{customer_id}")
async def api_update_customer(customer_id: int, data: dict = Body(...)):
    """更新客户"""
    customer = update_customer(customer_id, **data)
    if not customer:
        raise HTTPException(status_code=404, detail="客户不存在")
    return {"customer": customer}

@app.delete("/api/customers/{customer_id}")
async def api_delete_customer(customer_id: int):
    """删除客户"""
    delete_customer(customer_id)
    return {"status": "ok"}


@app.post("/api/customers/parse")
async def api_parse_customer(data: dict = Body(...)):
    """智能识别客户信息"""
    text = data.get("text", "")
    if not text:
        raise HTTPException(status_code=400, detail="文本不能为空")
    result = parse_customer_text(text)
    return {"parsed": result}


# ===== 消息管理 =====

@app.get("/api/customers/{customer_id}/messages")
async def api_get_messages(customer_id: int):
    """获取客户消息历史（自动拆分微信导出记录）"""
    customer = get_customer(customer_id)
    if not customer:
        raise HTTPException(status_code=404, detail="客户不存在")
    messages = get_messages(customer_id)
    
    # 调用消息解析器拆分微信导出记录
    if messages:
        messages = parse_and_split_messages(messages)
        # 按客户分组合并连续消息
        messages = group_consecutive_messages(messages)
    
    return {"messages": messages}


@app.post("/api/customers/{customer_id}/messages")
async def api_add_message(customer_id: int, data: dict = Body(...)):
    """添加消息"""
    customer = get_customer(customer_id)
    if not customer:
        raise HTTPException(status_code=404, detail="客户不存在")

    role = data.get("role", "user")
    content = data.get("content", "")
    timestamp = data.get("timestamp", "")
    session_id = data.get("session_id", "")

    if not content:
        raise HTTPException(status_code=400, detail="消息内容不能为空")

    msg = add_message(customer_id, role, content, timestamp, session_id)
    return {"message": msg}


@app.delete("/api/messages/{msg_id}")
async def api_delete_message(msg_id: int):
    """删除消息"""
    if not delete_message(msg_id):
        raise HTTPException(status_code=404, detail="消息不存在")
    return {"status": "ok"}


@app.put("/api/messages/{msg_id}")
async def api_update_message(msg_id: int, data: dict = Body(...)):
    """更新消息内容"""
    content = data.get("content", "")
    if not content:
        raise HTTPException(status_code=400, detail="内容不能为空")
    conn = get_db()
    conn.execute("UPDATE messages SET content=? WHERE id=?", (content, msg_id))
    conn.commit()
    conn.close()
    return {"status": "ok"}


# ===== 话术生成 =====

def _collect_effective_script_context(recent_text: str, customer: dict, analysis_tags: list, recent: str):
    """分析最近消息：学习有效话术 + 收集有效话术参考与反馈分析。

    返回 (effective_refs, feedback_analysis)
    """
    effective_refs = []
    feedback_analysis = ""
    try:
        if recent_text:
            import re
            # 按发送者分割消息块（不是按空行，因为话术多段）
            # 消息头模式：人名 + 日期时间（兼容 2026/08/11 与 08/11 两种格式）
            block_pattern = r'(?:^|\n)(?=[^\n]+ (?:\d{4}[/-])?\d{1,2}[/-]\d{1,2} \d{1,2}:\d{2})'
            raw_blocks = re.split(block_pattern, recent_text.strip())
            blocks = [b.strip() for b in raw_blocks if b.strip()]
            if len(blocks) >= 2:
                C_block = blocks[-1]
                B_block = blocks[-2]
                if "张兆渊" in B_block[:30] and "张兆渊" not in C_block[:30]:
                    B_lines = B_block.split("\n")
                    B_content = "\n".join(B_lines[1:]).strip() if len(B_lines) > 1 else B_block
                    C_lines = C_block.split("\n")
                    C_content = "\n".join(C_lines[1:]).strip() if len(C_lines) > 1 else C_block
                    if len(B_content) > 10:
                        score = 0
                        response_type = "中性"
                        scenario = "效果确认"
                        buy_words = ["买", "付款", "多少钱", "下单", "转账", "支付", "收款码", "怎么付"]
                        positive_words = ["好的", "瘦", "效果", "谢谢", "继续", "行", "可以", "不错", "有效", "轻了"]
                        reject_words = ["不", "别", "不用", "不要", "不需要", "算了", "再考虑", "太贵", "没钱", "自己来"]
                        for w in buy_words:
                            if w in C_content:
                                score += 3
                                response_type = "成交"
                                scenario = "复购推荐"
                                break
                        if response_type == "中性":
                            for w in positive_words:
                                if w in C_content:
                                    score += 2
                                    response_type = "积极"
                                    break
                        if response_type == "中性":
                            for w in reject_words:
                                if w in C_content:
                                    score -= 1
                                    response_type = "消极"
                                    break
                        if response_type == "中性" and len(C_content) > 0:
                            score += 1
                        if 30 <= len(B_content) <= 600:
                            score += 1
                        if "?" in B_content or "？" in B_content or "对不对" in B_content or "是吧" in B_content:
                            score += 1
                        if any(c.isdigit() for c in B_content):
                            score += 1
                        if "开心" in B_content or "高兴" in B_content or "放心" in B_content or "心疼" in B_content or "感动" in B_content:
                            score += 1
                        if "排油" in B_content or "排便" in B_content or "油脂" in B_content:
                            scenario = "排油跟进"
                        elif "瘦" in B_content or "体重" in B_content or "斤" in B_content:
                            scenario = "效果确认"
                        elif "买" in B_content or "续费" in B_content or "优惠" in B_content or "活动" in B_content:
                            scenario = "复购推荐"
                        elif "贵" in B_content or "价格" in B_content or "钱" in B_content:
                            scenario = "异议处理"
                        elif "阶段" in B_content or "周期" in B_content or "过程" in B_content:
                            scenario = "科普教育"
                        if score >= 2:
                            from effective_scripts import add_effective_script
                            ctype = customer.get("customer_type", "")
                            ctype_label = {"package": "套餐客户", "treatment": "疗程客户", "cid": "CID客户"}.get(ctype, "CID客户")
                            add_effective_script(B_content[:500], scenario, ctype_label, 1, score)
                        feedback_analysis = f"\n【话术效果分析】\n上次发送: {B_content[:100]}\n客户回应: {C_content[:100]}\n评分: {score} ({response_type})"
        try:
            from effective_scripts import search_effective_scripts_by_scenario
            scenario_keywords = analysis_tags[0] if analysis_tags else ""
            if not scenario_keywords:
                scenario_keywords = (recent or "")[:20]
            refs = search_effective_scripts_by_scenario(scenario_keywords, customer.get("customer_type", ""), top_k=3)
            if refs:
                effective_refs = [f"【有效话术参考】{r['content'][:60]}（评分{r['score']}，有效{r['effective_count']}次）" for r in refs]
        except Exception:
            pass
    except Exception:
        pass
    return effective_refs, feedback_analysis

@app.post("/api/generate")
async def api_generate_script(request: Request):
    """生成话术"""
    # 手动解析请求体
    try:
        data = await request.json()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"请求体解析失败: {str(e)}")
    customer_id = data.get("customer_id")
    recent = data.get("recent_messages", "")
    settings = data.get("settings", {})
    analysis_tags = data.get("tags", [])
    # 如果前端没传tags，从消息中自动提取#标签名
    if not analysis_tags:
        import re
        analysis_tags = re.findall(r'#([一-鿿\w]+)', recent)

    if not customer_id:
        raise HTTPException(status_code=400, detail="请选择客户")

    customer = get_customer(customer_id)
    if not customer:
        raise HTTPException(status_code=404, detail="客户不存在")

        # 获取聊天历史（同时包含客户消息和AI话术，最多60条）
    messages = get_messages(customer_id)
    chat_history_lines = []
    for m in messages[-60:]:
        if m['role'] == 'user':
            role_label = "客户"
        elif m['role'] == 'assistant':
            role_label = "我(销售顾问)"
        else:
            continue
        chat_history_lines.append(f"{m['timestamp']} {role_label}: {m['content']}")
    chat_history = "\n".join(chat_history_lines)

    # 搜索知识库
    search_data = search_knowledge(recent, top_k=20)
    knowledge_results = search_data.get("results", [])

    # 当前时间（优先用前端传的时间，否则用服务器时间）
    current_time = data.get("current_time", "") or datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    # 获取前端传的图片ID列表
    uploaded_image_ids = data.get("image_ids", [])

    # 处理图片相关逻辑
    enriched_recent = recent
    image_results = []

    # 获取该客户的所有图片
    customer_images = get_customer_images(customer_id)

    # 如果前端传了图片ID，则只使用这些图片（按上传顺序排列，老的图片不干扰当前上下文）
    if uploaded_image_ids:
        img_map = {img["id"]: img for img in customer_images}
        filtered = []
        for pid in uploaded_image_ids:
            if pid in img_map:
                filtered.append(img_map[pid])
        customer_images = filtered

    # ===== 匹配本地图片库（图片索引） =====
    local_image_matches = []
    if analysis_tags:
        # 用#标签作为搜索关键词
        local_image_matches = match_images_for_script(analysis_tags, top_k=5)
    elif uploaded_image_ids:
        # 有上传的图片但没标签，用图片描述作为关键词
        img_descs = []
        for img in customer_images[:3]:
            desc = img.get("description", "")
            if desc:
                img_descs.append(desc[:30])
        if img_descs:
            local_image_matches = match_images_for_script(img_descs, top_k=5)
    else:
        # 用最近消息前30字作为关键词模糊搜索
        search_kw = recent.strip()[:30] if recent.strip() else ""
        if search_kw:
            local_image_matches = match_images_for_script([search_kw], top_k=3)

    if customer_images:
        # 构建图片描述文本
        image_desc_lines = []
        for img in customer_images[:5]:
            desc = img.get("description", "")
            if desc:
                image_desc_lines.append(f"[图片: {img.get('original_name','')}] {desc}")
            else:
                image_desc_lines.append(f"[图片: {img.get('original_name','')}] (未识别)")

        # 情况1: 文本中包含 [图片] 标记 → 按顺序替换
        if "[图片]" in recent:
            parts = recent.split("[图片]")
            replaced_parts = [parts[0]]
            img_idx = 0
            for part in parts[1:]:
                if img_idx < len(customer_images):
                    desc = customer_images[img_idx].get("description", "")
                    if desc:
                        replaced_parts.append(f"[图片]({desc})")
                    else:
                        replaced_parts.append("[图片](未识别)")
                    img_idx += 1
                else:
                    replaced_parts.append("[图片](未匹配)")
                replaced_parts.append(part)
            enriched_recent = "".join(replaced_parts)

        # 情况2: 文本为空但有图片 → 用图片描述作为消息内容
        elif not recent.strip() and image_desc_lines:

            enriched_recent = "【图片描述】\n" + "\n".join(image_desc_lines)


        # 情况3: 文本有内容且有图片 → 在文本末尾附加图片参考
        elif recent.strip() and image_desc_lines:

            enriched_recent = recent + "\n\n【附带图片】\n" + "\n".join(image_desc_lines)


        # 构造 image_results 供 LLM 参考
        image_results = [{"id": img.get("id",""), "filename": img.get("filename",""),
                          "description": img.get("description",""),
                          "original_name": img.get("original_name","")}
                         for img in customer_images[:5]]

    # 去重标记
    recent_deduped = False

    # ===== 解析：分离聊天记录和销售顾问分析指令 =====
    def _parse_recent_messages(text: str):
        """从输入文本中分离聊天记录和销售顾问分析指令"""
        import re
        # 按空行分割成块，用 chr(10) 避免字符串跨行问题
        blocks = re.split(chr(10) + chr(10), text)
        chat_blocks = []
        instruction_blocks = []
        found_chat = False
        # 时间戳匹配：任意文本 + 数字/数字 + 数字:数字:数字
        ts_pat = re.compile(r'^.+?\s+\d{1,2}[/-]\d{1,2}\s+\d{1,2}:\d{2}(:\d{2})?\s*$')
        for block in blocks:
            block = block.strip()
            if not block:
                continue
            first_line = block.split(chr(10))[0].strip()
            if ts_pat.match(first_line):
                found_chat = True
                chat_blocks.append(block)
            elif found_chat:
                instruction_blocks.append(block)
            else:
                chat_blocks.append(block)
        return chr(10).join(chat_blocks), chr(10).join(instruction_blocks)

    # 调用 LLM 生成话术
    parsed_recent, parsed_instruction = _parse_recent_messages(enriched_recent)
    
    # 如果有标签，追加到最近消息中作为分析方向提示（直接追加到parsed_recent，避免被_parse_recent_messages过滤）
    if analysis_tags:
        tag_hint = "\n\n【本次分析方向】: " + "、".join(analysis_tags)
        parsed_recent = parsed_recent + tag_hint
    # ===== 有效话术分析 + 参考 =====
    recent_text = parsed_recent or recent or ""
    effective_refs, feedback_analysis = _collect_effective_script_context(recent_text, customer, analysis_tags, recent)

    # ===== 客户画像分析 =====
    customer_profile = None
    try:
        from customer_profile import analyze_customer_profile
        customer_profile = analyze_customer_profile(customer_id)
    except Exception as e:
        pass  # 分析失败不影响话术生成

    # ===== 调用LLM生成话术 =====
    result = await generate_script(
            customer_info=customer,
            recent_messages=parsed_recent,
            chat_history=chat_history,
            knowledge_results=knowledge_results,
            settings=settings,
            current_time=current_time,
            image_results=image_results,
            local_image_matches=local_image_matches,
            effective_refs=effective_refs,
            feedback_analysis=feedback_analysis
        )
    # 解析话术中的 [生成图片: 描述] 标记，不实际生成图片
    import re
    pending_images = re.findall(r'\[生成图片:\s*(.*?)\]', result)
    # 为每个待生成图片从本地图片库匹配相似图片
    pending_image_matches = {}
    for pi in pending_images:
        desc = pi.strip()
        if desc:
            # 用描述关键词搜索图片库
            sim_matches = search_images_hybrid(query=desc, limit=5)
            if sim_matches:
                pending_image_matches[desc] = sim_matches
    # 统计匹配图片的使用次数
    from database import increment_image_use_count
    for matches in pending_image_matches.values():
        for img in matches:
            if img.get('id'):
                try: increment_image_use_count(img['id'])
                except: pass
    if local_image_matches:
        for img in local_image_matches:
            if img.get('id'):
                try: increment_image_use_count(img['id'])
                except: pass
    return {
        'script': result,
        'pending_images': [{'description': d.strip()} for d in pending_images if d.strip()],
        'pending_image_matches': pending_image_matches if pending_image_matches else None,
        'degraded_reason': search_data.get('degraded_reason'),
        'local_image_matches': local_image_matches if local_image_matches else None,
        'customer_profile': customer_profile
        }


# ===== 流式生成话术（SSE） =====
@app.post("/api/generate/stream")
async def api_generate_script_stream(request: Request, data: dict = Body(...)):
    """流式生成话术，支持 SSE"""
    customer_id = data.get("customer_id")
    recent = data.get("recent_messages", "")
    settings = data.get("settings", {})
    analysis_tags = data.get("tags", [])
    if not analysis_tags:
        import re
        analysis_tags = re.findall(r'#([一-\u9fff\w]+)', recent)

    if not customer_id:
        return JSONResponse(status_code=400, content={"detail": "请选择客户"})

    customer = get_customer(customer_id)
    if not customer:
        return JSONResponse(status_code=404, content={"detail": "客户不存在"})

    messages = get_messages(customer_id)
    chat_history_lines = []
    for m in messages[-60:]:
        if m['role'] == 'user':
            role_label = "客户"
        elif m['role'] == 'assistant':
            role_label = "我(销售顾问)"
        else:
            continue
        chat_history_lines.append(f"{m['timestamp']} {role_label}: {m['content']}")
    chat_history = "\n".join(chat_history_lines)

    search_data = search_knowledge(recent, top_k=20)
    knowledge_results = search_data.get("results", [])
    current_time = data.get("current_time", "") or datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    uploaded_image_ids = data.get("image_ids", [])

    image_results = []
    customer_images = get_customer_images(customer_id)
    if uploaded_image_ids:
        img_map = {img["id"]: img for img in customer_images}
        filtered = []
        for pid in uploaded_image_ids:
            if pid in img_map:
                filtered.append(img_map[pid])
        customer_images = filtered

    local_image_matches = []
    enriched_recent = recent
    if analysis_tags:
        local_image_matches = match_images_for_script(analysis_tags, top_k=5)
    elif uploaded_image_ids:
        img_descs = []
        for img in customer_images[:3]:
            desc = img.get("description", "")
            if desc:
                img_descs.append(desc[:30])
        if img_descs:
            local_image_matches = match_images_for_script(img_descs, top_k=5)
    else:
        search_kw = recent.strip()[:30] if recent.strip() else ""
        if search_kw:
            local_image_matches = match_images_for_script([search_kw], top_k=3)

    if customer_images:
        image_desc_lines = []
        for img in customer_images[:5]:
            desc = img.get("description", "")
            if desc:
                image_desc_lines.append(f"[图片: {img.get('original_name','')}] {desc}")
            else:
                image_desc_lines.append(f"[图片: {img.get('original_name','')}] (未识别)")
        enriched_recent = recent
        if "[图片]" in recent:
            parts = recent.split("[图片]")
            replaced_parts = [parts[0]]
            img_idx = 0
            for part in parts[1:]:
                if img_idx < len(customer_images):
                    desc = customer_images[img_idx].get("description", "")
                    if desc:
                        replaced_parts.append(f"[图片]({desc})")
                    else:
                        replaced_parts.append("[图片](未识别)")
                    img_idx += 1
                else:
                    replaced_parts.append("[图片](未匹配)")
                replaced_parts.append(part)
            enriched_recent = "".join(replaced_parts)
        elif not recent.strip() and image_desc_lines:
            enriched_recent = "【图片描述】\n" + "\n".join(image_desc_lines)
        elif recent.strip() and image_desc_lines:
            enriched_recent = recent + "\n\n【附带图片】\n" + "\n".join(image_desc_lines)

        image_results = [{"id": img.get("id",""), "filename": img.get("filename",""),
                          "description": img.get("description",""),
                          "original_name": img.get("original_name","")}
                         for img in customer_images[:5]]

    def _parse_recent_messages(text: str):
        import re
        blocks = re.split(chr(10) + chr(10), text)
        chat_blocks = []
        instruction_blocks = []
        found_chat = False
        ts_pat = re.compile(r'^.+?\s+\d{1,2}[/-]\d{1,2}\s+\d{1,2}:\d{2}(:\d{2})?\s*$')
        for block in blocks:
            block = block.strip()
            if not block:
                continue
            first_line = block.split(chr(10))[0].strip()
            if ts_pat.match(first_line):
                found_chat = True
                chat_blocks.append(block)
            elif found_chat:
                instruction_blocks.append(block)
            else:
                chat_blocks.append(block)
        return chr(10).join(chat_blocks), chr(10).join(instruction_blocks)

    parsed_recent, parsed_instruction = _parse_recent_messages(enriched_recent)
    if analysis_tags:
        tag_hint = "\n\n【本次分析方向】: " + "、".join(analysis_tags)
        parsed_recent = parsed_recent + tag_hint

    degraded_reason = search_data.get('degraded_reason')

    # ===== 有效话术分析 + 参考 =====
    recent_text = parsed_recent or recent or ""
    effective_refs, feedback_analysis = _collect_effective_script_context(recent_text, customer, analysis_tags, recent)

    async def event_generator():
        # 先发 degraded_reason（如果有）
        if degraded_reason:
            yield f"data: {json.dumps({'type': 'degraded', 'content': degraded_reason})}\n\n"

        # 流式生成
        async for event in generate_script_stream(
            customer_info=customer,
            recent_messages=parsed_recent,
            chat_history=chat_history,
            knowledge_results=knowledge_results,
            settings=settings,
            current_time=current_time,
            image_results=image_results,
            local_image_matches=local_image_matches,
            effective_refs=effective_refs,
            feedback_analysis=feedback_analysis
        ):
            yield f"data: {json.dumps(event, ensure_ascii=False)}\n\n"

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        }
    )


    # ===== 搜索状态 =====

@app.get("/api/search/status")
async def api_search_status():
    """搜索状态"""
    from knowledge import is_embedding_available
    available = is_embedding_available()
    return {"embedding_available": available, "degraded_reason": None}


# ===== 设置管理 =====

@app.get("/api/settings")
async def api_get_settings():
    """获取设置"""
    config = cfg.load_config()
    return config.get("settings", {})


@app.put("/api/settings")
async def api_save_settings(data: dict = Body(...)):
    """保存设置"""
    config = cfg.load_config()
    config["settings"] = data
    cfg.save_config(config)
    return {"status": "ok"}


# ===== 个性化设置 =====

@app.get("/api/personalize")
async def api_get_personalize():
    """获取个性化设置"""
    config = cfg.load_config()
    return config.get("personalize", {"theme_color": "#0071e3", "bg_image": ""})


@app.put("/api/personalize")
async def api_save_personalize(data: dict = Body(...)):
    """保存个性化设置"""
    config = cfg.load_config()
    config["personalize"] = data
    cfg.save_config(config)
    return {"status": "ok"}


@app.post("/api/personalize/bg-image")
async def api_upload_bg_image(file: UploadFile = File(...)):
    """上传聊天背景图"""
    content = await file.read()
    save_dir = cfg.DATA_DIR / "bg_images"
    save_dir.mkdir(parents=True, exist_ok=True)
    ext = Path(file.filename or "bg.png").suffix.lower()
    save_name = f"chat_bg{ext}"
    save_path = save_dir / save_name
    save_path.write_bytes(content)
    # 保存路径到 config
    config = cfg.load_config()
    if "personalize" not in config:
        config["personalize"] = {}
    config["personalize"]["bg_image"] = str(save_path)
    cfg.save_config(config)
    return {"status": "ok", "path": str(save_path)}


@app.get("/api/personalize/bg-image")
async def api_get_bg_image():
    """获取聊天背景图"""
    config = cfg.load_config()
    bg_path = config.get("personalize", {}).get("bg_image", "")
    if bg_path and Path(bg_path).exists():
        return FileResponse(bg_path)
    raise HTTPException(status_code=404, detail="未设置背景图")


@app.delete("/api/personalize/bg-image")
async def api_delete_bg_image():
    """删除聊天背景图"""
    config = cfg.load_config()
    bg_path = config.get("personalize", {}).get("bg_image", "")
    if bg_path and Path(bg_path).exists():
        Path(bg_path).unlink()
    if "personalize" not in config:
        config["personalize"] = {}
    config["personalize"]["bg_image"] = ""
    cfg.save_config(config)
    return {"status": "ok"}


# ===== 系统提示词 =====

@app.get("/api/prompt")
async def api_get_prompt():
    """获取系统提示词"""
    prompt = cfg.get_system_prompt()
    return {"prompt": prompt}


@app.put("/api/prompt")
async def api_save_prompt(data: dict = Body(...)):
    """保存系统提示词"""
    prompt = data.get("prompt", "")
    cfg.save_system_prompt(prompt)
    from llm_client import reload_system_prompt
    reload_system_prompt()
    return {"status": "ok"}


# ===== 提示词管理（所有可编辑提示词） =====

@app.get("/api/settings/prompts")
async def api_get_all_prompts():
    """获取所有可编辑提示词"""
    return {"prompts": cfg.get_prompts()}


@app.post("/api/settings/prompts")
async def api_save_prompt_by_name(data: dict = Body(...)):
    """保存指定提示词"""
    name = data.get("name", "")
    content = data.get("content", "")
    if not name:
        raise HTTPException(status_code=400, detail="name 不能为空")
    cfg.save_prompt(name, content)
    # 如果保存的是系统提示词，重新加载
    if name == "system_prompt":
        from llm_client import reload_system_prompt
        reload_system_prompt()
    return {"status": "ok"}


# ===== 模型管理 =====

@app.get("/api/models/current")
async def api_get_current_model():
    """获取当前模型配置"""
    config = cfg.get_llm_config()
    masked = dict(config)
    key = masked.get("api_key", "")
    if len(key) > 8:
        masked["api_key"] = key[:8] + "****" + key[-4:]
    return {"model": masked}


@app.get("/api/models")
async def api_list_models():
    """获取可用模型列表（含三种默认ID）"""
    result = cfg.list_models()
    return result


@app.post("/api/models/test")
async def api_test_model(data: dict = Body(...)):
    """测试模型连接"""
    api_key = data.get("api_key", "")
    model = data.get("model", "")
    base_url = data.get("base_url", "")
    api_mode = data.get("api_mode", "openai")
    categories = data.get("categories", [])
    model_id = data.get("model_id", "")

    # 如果有 model_id，从配置取完整信息
    if model_id and not api_key:
        import config_manager as cfg
        config = cfg.load_config()
        m = config.get("models", {}).get("list", {}).get(model_id, {})
        if m:
            api_key = m.get("api_key", api_key)
            model = model or m.get("model", model)
            base_url = base_url or m.get("base_url", base_url)
            api_mode = api_mode or m.get("api_mode", "openai")

    is_embedding = "embedding" in categories or api_mode in ("dashscope_text", "dashscope_multimodal")
    is_imagegen = "imagegen" in categories

    if not api_key or not base_url or not model:
        return {"status": "error", "message": "参数不全"}

    try:
        import httpx
        async with httpx.AsyncClient(timeout=30) as client:
            if is_embedding:
                if api_mode == "dashscope_multimodal":
                    req_body = {"model": model, "input": {"contents": [{"text": "测试"}]}, "parameters": {"dimension": 256}}
                elif api_mode == "dashscope_text":
                    req_body = {"model": model, "input": {"texts": ["测试"]}}
                else:
                    req_body = {"input": "测试", "model": model, "dimensions": 1024}
                url = base_url.rstrip("/")
                if api_mode not in ("dashscope_multimodal", "dashscope_text") and not url.endswith("/embeddings"):
                    url += "/embeddings"
                resp = await client.post(url, headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}, json=req_body)
                if resp.status_code == 200:
                    data = resp.json()
                    if api_mode in ("dashscope_multimodal", "dashscope_text"):
                        vec = data.get("output", {}).get("embeddings", [{}])[0].get("embedding", [])
                    else:
                        vec = data.get("data", [{}])[0].get("embedding", [])
                    if vec:
                        return {"status": "ok", "message": f"✅ 连接成功（{len(vec)}维向量）"}
                    return {"status": "error", "message": "响应中未找到向量数据"}
                else:
                    return {"status": "error", "message": f"HTTP {resp.status_code}: {resp.text[:200]}"}
            elif is_imagegen:
                # 图片生成模型：发 /images/generations 测试
                url = base_url.rstrip("/")
                if not url.endswith("/images/generations"):
                    url += "/images/generations"
                resp = await client.post(url, headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
                    json={"model": model, "prompt": "测试", "n": 1, "size": "1024x1024"})
                if resp.status_code == 200:
                    return {"status": "ok", "message": "✅ 连接成功"}
                else:
                    return {"status": "error", "message": f"HTTP {resp.status_code}: {resp.text[:200]}"}
            else:
                url = base_url.rstrip("/")
                if not url.endswith("/chat/completions"):
                    url += "/chat/completions"
                resp = await client.post(url, headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
                    json={"model": model, "messages": [{"role": "user", "content": "Hi"}], "max_tokens": 5})
                if resp.status_code == 200:
                    return {"status": "ok", "message": "✅ 连接成功"}
                else:
                    return {"status": "error", "message": f"HTTP {resp.status_code}: {resp.text[:200]}"}
    except Exception as e:
        return {"status": "error", "message": str(e)[:200]}


@app.post("/api/models")
async def api_save_model(data: dict = Body(...)):
    """保存模型配置"""
    provider = data.get("provider", "")
    api_key = data.get("api_key", "")
    model = data.get("model", "")
    base_url = data.get("base_url", "")
    set_default = data.get("set_default", False)
    categories = data.get("categories", [])
    api_mode = data.get("api_mode", "openai")
    if not api_key or not base_url:
        raise HTTPException(status_code=400, detail="API Key 和 Base URL 不能为空")
    result = cfg.add_model(provider, api_key, model, base_url, categories, api_mode)
    if set_default:
        cfg.set_default_model(result["id"])
    return {"status": "ok", "model": result}


@app.delete("/api/models/{model_id}")
async def api_delete_model(model_id: str):
    """删除模型配置"""
    try:
        cfg.remove_model(model_id)
        return {"status": "ok"}
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.put("/api/models/{model_id}/default")
async def api_set_default_model(model_id: str):
    """将已有模型设为默认（话术模型）"""
    cfg.set_default_model(model_id)
    return {"status": "ok"}


@app.put("/api/models/{model_id}/vision-default")
async def api_set_vision_default_model(model_id: str):
    """设置图片识别模型"""
    cfg.set_vision_default_model(model_id)
    return {"status": "ok"}


@app.put("/api/models/{model_id}/image-gen-default")
async def api_set_image_gen_default_model(model_id: str):
    """设置图片生成模型"""
    cfg.set_image_gen_default_model(model_id)
    return {"status": "ok"}


@app.put("/api/models/{model_id}/embedding-default")
async def api_set_embedding_default_model(model_id: str):
    """设置向量模型"""
    cfg.set_embedding_default_model(model_id)
    return {"status": "ok"}


# ===== 知识库管理 =====

@app.get("/api/knowledge/documents")
async def api_list_documents():
    """列出知识库文档"""
    docs = list_documents()
    return {"documents": docs}


@app.post("/api/knowledge/upload")
async def api_upload_document(file: UploadFile = File(...)):
    """上传知识库文件"""
    import tempfile
    content_bytes = await file.read()
    suffix = Path(file.filename).suffix if file.filename else ".txt"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(content_bytes)
        tmp_path = tmp.name
    try:
        text = extract_text_from_file(tmp_path)
        if not text.strip():
            raise HTTPException(status_code=400, detail="无法提取文件内容")
        doc = add_document(file.filename, text)
        return {"status": "ok", "document": doc}
    finally:
        try:
            os.unlink(tmp_path)
        except:
            pass


@app.post("/api/knowledge/reindex")
async def api_reindex_knowledge():
    """重新索引知识库"""
    reindex_all()
    return {"status": "ok"}


@app.delete("/api/knowledge/documents/{doc_id}")
async def api_delete_document(doc_id: str):
    """删除知识库文档"""
    if not delete_document(doc_id):
        raise HTTPException(status_code=404, detail="文档不存在")
    return {"status": "ok"}

@app.get("/api/knowledge/documents/{doc_id}/chunks")
async def api_get_document_chunks(doc_id: str):
    """获取文档的片段列表（含向量预览）"""
    try:
        import sqlite3
        from knowledge import DB_PATH, VEC_DB_PATH, _get_vec_conn
        # 从向量库获取片段（knowledge_chunks 在向量库）
        mconn = _get_vec_conn()
        try:
            rows = mconn.execute("SELECT id, chunk_index, content FROM knowledge_chunks WHERE doc_id = ? ORDER BY chunk_index", (doc_id,)).fetchall()
        except Exception:
            rows = []
        mconn.close()
        # 从向量库获取向量预览
        vec_map = {}
        try:
            vconn = _get_vec_conn()
            vec_rows = vconn.execute("SELECT id, vector FROM knowledge_vectors WHERE id LIKE ? ORDER BY id", (doc_id + '_%',)).fetchall()
            for vr in vec_rows:
                vid = vr['id']
                vec_bytes = vr['vector'] if vr['vector'] else b''
                if isinstance(vec_bytes, bytes) and len(vec_bytes) > 0:
                    dim = len(vec_bytes) // 4
                    preview = '[' + ','.join(str(round(b/255,3)) for b in vec_bytes[:12]) + '...]' if len(vec_bytes) >= 4 else '(空)'
                    vec_map[vid] = {'vector_dim': dim, 'vector_preview': preview}
                else:
                    vec_map[vid] = {'vector_dim': 0, 'vector_preview': '(空)'}
            vconn.close()
        except Exception:
            pass
        chunks = []
        for row in rows:
            vi = vec_map.get(row['id'], {})
            # 从文本中提取前几个词作为关键词
            content = row['content'] or ''
            # 简单关键词提取：取前 3-5 个有意义的词
            words = [w for w in content.replace('\n', ' ').split() if len(w) > 1][:5]
            keywords = words if words else []
            chunks.append({
                'id': row['id'],
                'content': content,
                'chunk_index': row['chunk_index'],
                'keywords': keywords,
                'vector_dim': vi.get('vector_dim', 0),
                'vector_preview': vi.get('vector_preview', '')
            })
        return {"doc_id": doc_id, "chunks": chunks, "total_chunks": len(chunks)}
    except Exception as e:
        return {"doc_id": doc_id, "chunks": [], "total_chunks": 0, "error": str(e)}


@app.get("/api/knowledge/documents/{doc_id}/file")
async def api_get_document_file(doc_id: str):
    """下载知识库源文件"""
    from pathlib import Path as P
    from fastapi.responses import FileResponse
    from knowledge import KNOWLEDGE_DIR, _load_meta
    meta = _load_meta()
    for d in meta:
        if d["id"] == doc_id:
            fp = KNOWLEDGE_DIR / d["filename"]
            if fp.exists():
                return FileResponse(str(fp), filename=d["filename"])
            raise HTTPException(status_code=404, detail="文件不存在")
    raise HTTPException(status_code=404, detail="文档不存在")



@app.get("/api/knowledge/search")
async def api_search_knowledge(q: str = Query(..., min_length=1)):
    """搜索知识库"""
    return search_knowledge(q, top_k=20)


# ===== 客户图片上传 =====


@app.get("/api/customers/{customer_id}/images")
async def api_get_customer_images(customer_id: int):
    """获取客户图片列表"""
    images = get_customer_images(customer_id)
    return {"images": images}


@app.post("/api/customers/{customer_id}/images")
async def api_upload_customer_image(customer_id: int, file: UploadFile = File(...), description: str = ""):
    """上传客户图片"""
    content_bytes = await file.read()
    img = add_customer_image(customer_id, file.filename or "image.png", content_bytes, description)
    return {"status": "ok", "image": img}


@app.get("/api/customers/images/{img_id}/file")
async def api_get_image_file(img_id: str):
    """获取图片文件"""
    from database import get_db
    conn = get_db()
    cursor = conn.execute("SELECT * FROM customer_images WHERE id=?", (img_id,))
    row = cursor.fetchone()
    conn.close()
    if not row:
        raise HTTPException(status_code=404, detail="图片不存在")
    img = dict(row)
    img_path = Path(img.get("filepath", ""))
    if not img_path.is_absolute():
        img_path = cfg.DATA_DIR / "images" / img_path
    if not img_path.exists():
        raise HTTPException(status_code=404, detail="图片文件不存在")
    return FileResponse(str(img_path), media_type="image/png", filename=img.get("original_name", "image.png"))


@app.get("/api/image-index/images/{img_id}/file")
async def api_get_image_index_file(img_id: int):
    """获取图片索引库中的图片文件"""
    from database import get_db
    conn = get_db()
    cursor = conn.execute("SELECT file_path FROM image_index WHERE id=?", (img_id,))
    row = cursor.fetchone()
    conn.close()
    if not row or not row["file_path"]:
        # 返回占位图而非404，避免前端控制台报错
        return Response(content=PLACEHOLDER_SVG, media_type="image/svg+xml")
    img_path = Path(row["file_path"])
    if not img_path.exists():
        # 文件已被移动或删除，返回占位图
        return Response(content=PLACEHOLDER_SVG, media_type="image/svg+xml")
    return FileResponse(str(img_path))


@app.delete("/api/customers/images/{img_id}")
async def api_delete_customer_image(img_id: str):
    """删除客户图片"""
    if not delete_customer_image(img_id):
        raise HTTPException(status_code=404, detail="图片不存在")
    return {"status": "ok"}


# ===== 生成图片服务 =====

@app.get("/api/generated_images/{filename}")
async def serve_generated_image(filename: str):
    """提供AI生成的图片文件"""
    from pathlib import Path
    from fastapi.responses import FileResponse, Response
    img_path = GENERATED_IMAGES_DIR / filename
    if not img_path.exists():
        # 返回占位图而非404，避免前端控制台报错
        return Response(content=PLACEHOLDER_SVG, media_type="image/svg+xml")
    return FileResponse(str(img_path))

# ===== 标签管理（存储在 config.json 的 tags 键下） =====
def _load_tags():
    """加载标签列表"""
    from config_manager import load_config
    config = load_config()
    return config.get("tags", [])

def _save_tags(tags: list):
    """保存标签列表"""
    from config_manager import load_config, save_config
    config = load_config()
    config["tags"] = tags
    save_config(config)

@app.get("/api/tags")
async def api_get_tags():
    """获取标签列表"""
    return {"tags": _load_tags()}

@app.post("/api/tags")
async def api_add_tag(data: dict = Body(...)):
    """添加标签"""
    tag = data.get("tag", "").strip()
    if not tag:
        raise HTTPException(status_code=400, detail="标签名不能为空")
    tags = _load_tags()
    if tag in [t["name"] for t in tags]:
        raise HTTPException(status_code=400, detail="标签已存在")
    tags.append({"id": str(len(tags) + 1), "name": tag})
    _save_tags(tags)
    return {"status": "ok", "tags": tags}

@app.delete("/api/tags/{tag_id}")
async def api_delete_tag(tag_id: str):
    """删除标签"""
    tags = _load_tags()
    tags = [t for t in tags if t["id"] != tag_id]
    _save_tags(tags)
    return {"status": "ok", "tags": tags}

@app.put("/api/tags/{tag_id}")
async def api_update_tag(tag_id: str, data: dict = Body(...)):
    """修改标签名"""
    tag_name = data.get("name", "").strip()
    if not tag_name:
        raise HTTPException(status_code=400, detail="标签名不能为空")
    tags = _load_tags()
    for t in tags:
        if t["id"] == tag_id:
            t["name"] = tag_name
            break
    _save_tags(tags)
    return {"status": "ok", "tags": tags}

# ===== 数据导出 =====

@app.get("/api/export")
async def api_export_data():
    """导出所有用户数据"""
    from config_manager import load_config, get_system_prompt
    from knowledge import list_documents
    import io, json
    from fastapi.responses import JSONResponse
    
    # 收集所有客户数据
    all_customers = list_customers()
    all_messages = []
    for c in all_customers:
        cid = c.get("id")
        msgs = get_messages(cid)
        for m in msgs:
            m["customer_id"] = cid
        all_messages.extend(msgs)
    
    config = load_config()
    system_prompt = get_system_prompt()
    banned_words = config.get('settings', {}).get('banned_words', '')
    
    export_data = {
        "customers": all_customers,
        "messages": all_messages,
        "system_prompt": system_prompt,
        "banned_words": config.get("settings", {}).get("banned_words", ""),
        "settings": config.get("settings", {}),
        "tags": config.get("tags", []),
    }
    
    return JSONResponse(content=export_data)

# ===== 图片索引 API =====

@app.get("/api/image-index/status")
async def api_image_index_status():
    """获取图片索引统计"""
    return {"status": "ok", "data": get_image_index_status()}


@app.get("/api/image-index/progress")
async def api_image_index_progress():
    """获取当前任务进度"""
    return {"status": "ok", "data": dict(_scan_progress)}

@app.post("/api/image-index/scan")
async def api_image_index_scan(data: dict = Body(...)):
    """扫描文件夹并更新进度"""
    from pathlib import Path
    base_path = data.get("path", "")
    if not base_path:
        raise HTTPException(status_code=400, detail="请选择文件夹路径")
    
    _scan_progress["total"] = 0
    _scan_progress["pending"] = 0
    _scan_progress["done"] = 0
    _scan_progress["failed"] = 0
    _scan_progress["skipped"] = 0
    _scan_progress["current"] = None
    _scan_progress["status"] = "scanning"
    _scan_progress["cancel_flag"] = False
    _scan_progress["pause_flag"] = False
    _scan_progress["task_type"] = "scan"
    _scan_progress["base_path"] = base_path
    import time
    _scan_progress["started_at"] = time.strftime("%Y-%m-%d %H:%M:%S")
    _scan_progress["completed_at"] = None
    
    scan_results = scan_folder_for_index(base_path)
    if not scan_results:
        _scan_progress["status"] = "error"
        _scan_progress["error_msg"] = "文件夹不存在或无图片文件"
        return {"status": "error", "message": "文件夹不存在或无图片文件，请检查路径是否正确"}
    _scan_progress["total"] = len(scan_results)
    
    sync_result = image_index_sync(scan_results)
    _scan_progress["status"] = "ready"
    _scan_progress["pending"] = len(sync_result.get("new", [])) + len(sync_result.get("reindex", []))
    _scan_progress["skipped"] = sync_result.get("existing", 0)
    _scan_progress["base_path"] = base_path
    
    # 按分类统计
    from database import get_db
    _conn = get_db()
    _cat_rows = _conn.execute("SELECT category, COUNT(*) as cnt FROM image_index WHERE status IN ('pending','done','failed') GROUP BY category ORDER BY category").fetchall()
    _conn.close()
    _scan_progress["categories"] = [{"name": r["category"], "total": r["cnt"], "done": 0, "failed": 0, "status": "waiting"} for r in _cat_rows]
    _scan_progress["total_categories"] = len(_scan_progress["categories"])
    _scan_progress["done_categories"] = 0
    _scan_progress["current_category"] = ""
    
    return {"status": "ok", "data": {"scan": sync_result, "progress": _scan_progress.copy()}}

@app.post("/api/image-index/scan-incremental")
async def api_image_index_scan_incremental(data: dict = Body(...)):
    """
    增量扫描：传入路径和分类名。
    1. 检查分类是否存在，不存在则自动创建
    2. 扫描文件夹中的图片
    3. 只新增未在数据库中的图片，不删除任何已有图片
    4. 自动触发批量识别
    """
    category_name = data.get("category_name", "").strip()
    base_path = data.get("path", "").strip()
    if not base_path or not category_name:
        return {"status": "error", "message": "缺少路径或分类名称"}

    from pathlib import Path
    path_obj = Path(base_path)
    if not path_obj.exists() or not path_obj.is_dir():
        return {"status": "error", "message": f"文件夹不存在: {base_path}"}

    # 扫描文件夹
    scan_results = scan_folder_for_index(base_path)
    if not scan_results:
        return {"status": "ok", "data": {"new_count": 0, "skipped": 0, "message": "文件夹中未找到图片"}}

    # 增量同步（不删图）
    result = image_index_scan_incremental(category_name, scan_results)

    # 如果有新图片，自动启动批量识别
    if result["new_count"] > 0:
        global _recognize_task
        current_status = _scan_progress.get("status", "idle")
        if current_status not in ("recognizing", "scanning", "starting"):
            _scan_progress["total"] = 0
            _scan_progress["pending"] = 0
            _scan_progress["done"] = 0
            _scan_progress["failed"] = 0
            _scan_progress["skipped"] = 0
            _scan_progress["current"] = None
            _scan_progress["status"] = "starting"
            _scan_progress["cancel_flag"] = False
            _scan_progress["pause_flag"] = False
            _scan_progress["cancel_categories"] = []
            _scan_progress["task_type"] = "recognize"
            _scan_progress["base_path"] = base_path
            _scan_progress["started_at"] = __import__("time").strftime("%Y-%m-%d %H:%M:%S")
            _scan_progress["completed_at"] = None
            _scan_progress["current_category"] = ""
            _scan_progress["done_categories"] = 0
            _scan_progress["categories"] = []
            asyncio.ensure_future(_safe_cancel_task()).add_done_callback(lambda _: _start_recognition())

    return {
        "status": "ok",
        "data": {
            "new_count": result["new_count"],
            "skipped": result["skipped"],
            "message": f"新增 {result['new_count']} 张图片，跳过 {result['skipped']} 张已有图片"
        }
    }

@app.post("/api/image-index/init")
async def api_image_index_init():
    """清空图片索引"""
    image_index_clear()
    return {"status": "ok"}

@app.get("/api/image-index/pending")
async def api_image_index_pending(limit: int = 50):
    """获取待识别图片"""
    return {"status": "ok", "data": get_pending_images(limit)}

@app.post("/api/image-index/recognize")
async def api_image_index_recognize(data: dict = Body(...)):
    """识别单张图片"""
    image_id = data.get("id")
    file_path = data.get("path", "")
    if not file_path:
        raise HTTPException(status_code=400, detail="缺少文件路径")
    try:
        from llm_client import _analyze_image
        desc, tags = await _analyze_image(file_path)
        update_image_result(image_id, desc, tags)
        return {"status": "ok", "data": {"description": desc, "tags": tags}}
    except Exception as e:
        error_msg = str(e)
        update_image_result(image_id, "", [], error_msg=error_msg)
        return {"status": "error", "data": {"error": error_msg}}


# ===== 批量识别 =====

_recognize_task = None

async def _safe_cancel_task():
    global _recognize_task
    if _recognize_task and not _recognize_task.done():
        _recognize_task.cancel()
        try:
            await _recognize_task
        except:
            pass
    _recognize_task = None

def _start_recognition():
    global _recognize_task
    _scan_progress["total"] = 0
    _scan_progress["pending"] = 0
    _scan_progress["done"] = 0
    _scan_progress["failed"] = 0
    _scan_progress["skipped"] = 0
    _scan_progress["status"] = "starting"
    _scan_progress["cancel_flag"] = False
    _scan_progress["pause_flag"] = False
    _scan_progress["task_type"] = "recognize"
    _recognize_task = asyncio.create_task(_run_recognize_batch())

async def _run_recognize_batch():
    """后台运行批量识别（轮询模型）"""
    from llm_client import _analyze_image_with_model
    from config_manager import get_vision_model_list
    import time
    from collections import defaultdict

    # 获取所有vision模型（轮询池）
    vision_models = get_vision_model_list()
    if not vision_models:
        from config_manager import get_vision_model_config
        vision_models = [get_vision_model_config()]
    model_idx = 0
    
    # 获取所有待识别图片，按分类分组
    pending = get_pending_images(limit=10000)
    cat_groups = defaultdict(list)
    for img in pending:
        cat_groups[img.get("category", "未分类")].append(img)
    sorted_cats = sorted(cat_groups.keys())
    total = len(pending)
    total_cats = len(sorted_cats)
    
    _scan_progress["total"] = total
    _scan_progress["pending"] = total
    _scan_progress["total_categories"] = total_cats
    _scan_progress["categories"] = [
        {"name": cat, "total": len(cat_groups[cat]), "done": 0, "failed": 0, "status": "waiting"}
        for cat in sorted_cats
    ]
    _scan_progress["status"] = "recognizing"
    _scan_progress["started_at"] = time.strftime("%Y-%m-%d %H:%M:%S")
    _scan_progress["completed_at"] = None
    _scan_progress["current_category"] = ""
    _scan_progress["done_categories"] = 0
    _scan_progress["current"] = None
    
    if total == 0:
        _scan_progress["status"] = "idle"
        _scan_progress["completed_at"] = time.strftime("%Y-%m-%d %H:%M:%S")
        return
    
    for cat_idx, cat_name in enumerate(sorted_cats):
        if _scan_progress["cancel_flag"]:
            break
        
        # 检查该分类是否被用户取消
        if cat_name in _scan_progress.get("cancel_categories", []):
            # 跳过该分类：图片保持 pending 状态
            for c in _scan_progress["categories"]:
                if c["name"] == cat_name:
                    c["done"] = 0
                    c["failed"] = 0
                    c["status"] = "cancelled"
                    break
            _scan_progress["done_categories"] = cat_idx + 1
            continue
        
        _scan_progress["current_category"] = cat_name
        _scan_progress["done_categories"] = cat_idx
        
        for c in _scan_progress["categories"]:
            if c["name"] == cat_name:
                c["status"] = "processing"
                break
        
        cat_images = cat_groups[cat_name]
        cat_done = 0
        cat_failed = 0
        
        for img in cat_images:
            if _scan_progress["cancel_flag"]:
                _scan_progress["status"] = "cancelled"
                _scan_progress["completed_at"] = time.strftime("%Y-%m-%d %H:%M:%S")
                break
            
            # 检查该分类是否在循环中被取消
            if cat_name in _scan_progress.get("cancel_categories", []):
                remaining = cat_images[cat_images.index(img):]
                for c in _scan_progress["categories"]:
                    if c["name"] == cat_name:
                        c["done"] = cat_done
                        c["failed"] = cat_failed
                        c["status"] = "cancelled"
                        break
                # 取消后剩余图片回退到 pending（不计数）
                _scan_progress["pending"] -= len(remaining)
                break
            
            while _scan_progress["pause_flag"] and not _scan_progress["cancel_flag"]:
                await asyncio.sleep(0.5)
            
            _scan_progress["current"] = {
                "id": img["id"],
                "file_path": img["file_path"],
                "filename": str(img["file_path"]).split("/")[-1].split("\\")[-1],
                "description": ""
            }
            _scan_progress["pending"] -= 1
            
            try:
                from pathlib import Path
                if not Path(img["file_path"]).exists():
                    update_image_result(img["id"], "", [], error_msg="文件不存在")
                    _scan_progress["failed"] += 1
                    cat_failed += 1
                    continue
                
                desc, tags, customers = None, [], ""
                import time
                now = time.time()
                # 轮询选择模型（按失败率动态跳过）
                for attempt in range(len(vision_models)):
                    current_model = vision_models[model_idx % len(vision_models)]
                    model_idx += 1
                    mid = current_model.get("_id", "")
                    stats = _model_stats.setdefault(mid, {"success":0, "fail":0, "consecutive_fail":0, "last_try":0})
                    # 连续失败3次以上，跳过60秒
                    if stats["consecutive_fail"] >= 3 and now - stats["last_try"] < 60:
                        continue
                    model_label = mid or current_model.get("model", "?")
                    print(f"  [识别] 图{img['id']} → {model_label}")
                    try:
                        desc, tags, customers = await _analyze_image_with_model(
                            img["file_path"], cat_name, current_model)
                        stats["success"] += 1
                        stats["consecutive_fail"] = max(0, stats["consecutive_fail"] - 1)
                        stats["last_try"] = now
                        break  # 成功
                    except Exception as e:
                        err_msg = str(e)[:80]
                        stats["fail"] += 1
                        stats["consecutive_fail"] += 1
                        stats["last_try"] = now
                        if attempt < len(vision_models) - 1:
                            continue  # 试下一个模型
                        # 所有模型都失败了
                        raise Exception(f"所有模型均失败: {err_msg}")
                update_image_result(img["id"], desc, tags, applicable_customers=customers)
                _scan_progress["done"] += 1
                cat_done += 1
                _scan_progress["current"]["description"] = desc[:80]
            except Exception as e:
                error_msg = str(e)
                update_image_result(img["id"], "", [], error_msg=error_msg)
                _scan_progress["failed"] += 1
                cat_failed += 1
            
            for c in _scan_progress["categories"]:
                if c["name"] == cat_name:
                    c["done"] = cat_done
                    c["failed"] = cat_failed
                    break
        
        # [限流保护] 每张图间隔6秒，约10张/分钟
        await asyncio.sleep(6)
        
        # [限流保护] 大分类每20张额外休息
        if cat_done > 0 and cat_done % 20 == 0:
            await asyncio.sleep(15)
        
        if cat_name in _scan_progress.get("cancel_categories", []):
            # 已被取消的分类保持 cancelled 状态
            for c in _scan_progress["categories"]:
                if c["name"] == cat_name:
                    c["status"] = "cancelled"
                    break
        elif not _scan_progress["cancel_flag"]:
            for c in _scan_progress["categories"]:
                if c["name"] == cat_name:
                    c["status"] = "done"
                    break
            _scan_progress["done_categories"] = cat_idx + 1
        
        # [限流保护] 每个分类完成后休息30秒
        if not _scan_progress["cancel_flag"] and cat_name not in _scan_progress.get("cancel_categories", []):
            await asyncio.sleep(30)
            # 每10个分类多休息2分钟
            if (cat_idx + 1) % 10 == 0:
                await asyncio.sleep(120)
    
    if _scan_progress["cancel_flag"]:
        _scan_progress["status"] = "cancelled"
    else:
        _scan_progress["status"] = "idle"
        _scan_progress["current"] = None
        _scan_progress["current_category"] = ""
    _scan_progress["completed_at"] = time.strftime("%Y-%m-%d %H:%M:%S")


@app.post("/api/image-index/recognize-batch")
async def api_image_index_recognize_batch(data: dict = Body(...)):
    """批量识别所有待处理图片（后台任务，立即返回）"""
    global _recognize_task
    import time
    
    _scan_progress["total"] = 0
    _scan_progress["pending"] = 0
    _scan_progress["done"] = 0
    _scan_progress["failed"] = 0
    _scan_progress["skipped"] = 0
    _scan_progress["current"] = None
    _scan_progress["status"] = "starting"
    _scan_progress["cancel_flag"] = False
    _scan_progress["pause_flag"] = False
    _scan_progress["cancel_categories"] = []
    _scan_progress["task_type"] = "recognize"
    _scan_progress["base_path"] = data.get("base_path", _scan_progress.get("base_path", ""))
    _scan_progress["started_at"] = time.strftime("%Y-%m-%d %H:%M:%S")
    _scan_progress["completed_at"] = None
    _scan_progress["current_category"] = ""
    _scan_progress["done_categories"] = 0
    _scan_progress["categories"] = []
    
    # 启动后台任务
    asyncio.ensure_future(_safe_cancel_task()).add_done_callback(lambda _: _start_recognition())
    
    return {"status": "started", "data": {"message": "批量识别已启动"}}


@app.get("/api/image-index/search")
async def api_image_index_search(q: str = "", category: str = "", limit: int = 20):
    """搜索图片"""
    return {"status": "ok", "data": search_images_hybrid(query=q, category=category, limit=limit)}


# 图片重建向量索引进度
_image_reindex_progress = {"running": False, "total": 0, "done": 0, "error": None, "message": "", "failed_errors": {}}


@app.post("/api/image-index/reindex")
async def api_image_index_reindex():
    """重建图片向量索引（后台运行）"""
    if _image_reindex_progress["running"]:
        return {"status": "error", "message": "正在重建中，请稍候"}
    
    # 后台启动
    import threading
    def _run():
        try:
            _image_reindex_progress["running"] = True
            _image_reindex_progress["error"] = None
            _image_reindex_progress["message"] = "正在生成向量..."
            _image_reindex_progress["failed_errors"] = {}
            # 设置进度回调
            def _progress(done, total, success, failed, failed_errors=None):
                            _image_reindex_progress["done"] = done
                            _image_reindex_progress["total"] = total
                            _image_reindex_progress["message"] = f"向量生成中… {done}/{total}（成功{success}，失败{failed}）"
                            _image_reindex_progress["failed_errors"] = failed_errors or {}
            import database
            database._image_reindex_progress_callback = _progress
            result = rebuild_image_vectors()
            if isinstance(result, dict):
                cnt = result.get("success", 0)
                _image_reindex_progress["failed_errors"] = result.get("failed_errors", {})
            else:
                cnt = result
            _image_reindex_progress["done"] = cnt
            err_count = len(_image_reindex_progress.get("failed_errors", {}))
            if err_count > 0:
                _image_reindex_progress["message"] = f"重建完成，成功 {cnt} 张，失败 {err_count} 张"
            else:
                _image_reindex_progress["message"] = f"重建完成，成功 {cnt} 张"
        except Exception as e:
            _image_reindex_progress["error"] = str(e)
            _image_reindex_progress["message"] = f"重建失败: {e}"
        finally:
            _image_reindex_progress["running"] = False
            import database
            database._image_reindex_progress_callback = None
    t = threading.Thread(target=_run, daemon=True)
    t.start()
    return {"status": "ok", "data": {"message": "向量索引重建已启动，请稍候查看进度"}}


@app.get("/api/image-index/reindex/progress")
async def api_image_index_reindex_progress():
    """获取重建向量索引的进度"""
    return {"status": "ok", "data": _image_reindex_progress.copy()}

@app.post("/api/image-index/reindex/retry-failed")
async def api_image_index_reindex_retry_failed():
    """重试重建失败的图片向量"""
    if _image_reindex_progress.get("running"):
        return {"status": "error", "message": "正在重建中，请稍候"}
    errs = _image_reindex_progress.get("failed_errors", {})
    if not errs:
        return {"status": "ok", "data": {"message": "没有失败的图片"}}
    # 重置进度
    _image_reindex_progress["running"] = True
    _image_reindex_progress["error"] = None
    _image_reindex_progress["message"] = "正在重试失败的图片..."
    _image_reindex_progress["failed_errors"] = {}
    import threading
    def _run():
        try:
            import database
            def _progress(done, total, success, failed, failed_errors=None):
                            _image_reindex_progress["done"] = done
                            _image_reindex_progress["total"] = total
                            _image_reindex_progress["message"] = f"向量重试中… {done}/{total}（成功{success}，失败{failed}）"
                            _image_reindex_progress["failed_errors"] = failed_errors or {}
            database._image_reindex_progress_callback = _progress
            result = database.rebuild_image_vectors()
            if isinstance(result, dict):
                cnt = result.get("success", 0)
            else:
                cnt = result
            _image_reindex_progress["done"] = cnt
            _image_reindex_progress["message"] = f"重试完成，成功 {cnt} 张"
        except Exception as e:
            _image_reindex_progress["error"] = str(e)
            _image_reindex_progress["message"] = f"重试失败: {e}"
        finally:
            _image_reindex_progress["running"] = False
            import database
            database._image_reindex_progress_callback = None
    t = threading.Thread(target=_run, daemon=True)
    t.start()
    return {"status": "ok", "data": {"message": "重试已启动"}}




# ===== 图片索引进度 API =====

@app.get("/api/image-index/progress")
async def api_image_index_progress():
    """获取当前处理进度"""
    return {"status": "ok", "data": _scan_progress.copy()}


@app.post("/api/image-index/progress/pause")
async def api_image_index_pause():
    """暂停处理"""
    _scan_progress["pause_flag"] = True
    return {"status": "ok"}


@app.post("/api/image-index/progress/resume")
async def api_image_index_resume():
    """恢复处理"""
    _scan_progress["pause_flag"] = False
    return {"status": "ok"}


@app.post("/api/image-index/progress/cancel")
async def api_image_index_cancel():
    """取消处理"""
    _scan_progress["cancel_flag"] = True
    _scan_progress["status"] = "cancelled"
    return {"status": "ok"}


@app.post("/api/image-index/progress/cancel-category")
async def api_image_index_cancel_category(data: dict = Body(...)):
    """取消单个分类"""
    cat_name = data.get("name", "")
    if not cat_name:
        raise HTTPException(status_code=400, detail="请提供分类名称")
    if "cancel_categories" not in _scan_progress:
        _scan_progress["cancel_categories"] = []
    if cat_name not in _scan_progress["cancel_categories"]:
        _scan_progress["cancel_categories"].append(cat_name)
    return {"status": "ok", "data": {"cancelled": cat_name}}


@app.post("/api/image-index/delete-category")
async def api_image_index_delete_category(data: dict = Body(...)):
    """删除单个分类的所有图片"""
    cat_name = data.get("name", "")
    if not cat_name:
        raise HTTPException(status_code=400, detail="请提供分类名称")
    from database import delete_image_category
    deleted = delete_image_category(cat_name)
    # 也从 _scan_progress 的 categories 列表中移除
    if "categories" in _scan_progress and _scan_progress["categories"]:
        _scan_progress["categories"] = [c for c in _scan_progress["categories"] if c["name"] != cat_name]
        _scan_progress["total_categories"] = len(_scan_progress["categories"])
    return {"status": "ok", "data": {"deleted": deleted, "category": cat_name}}




@app.post("/api/image-index/folder")
async def api_image_index_folder(data: dict = Body(...)):
    """选择文件夹：按文件夹名自动创建/匹配子分类，并索引未识别图片"""
    name = data.get("name", "")
    files = data.get("files", [])
    if not name or not files:
        raise HTTPException(status_code=400, detail="name 和 files 不能为空")
    from database import add_image_record, get_all_subcategories, create_subcategory
    # 确保子分类存在
    subs = get_all_subcategories()
    if name not in [s["name"] for s in subs]:
        create_subcategory(name)
    # 逐文件处理：已存在则更新状态，不存在则写入 pending
    added = 0
    for f in files:
        path = f.get("path", "")
        if not path:
            continue
        img_id = add_image_record(path, [name], status="pending")
        if img_id:
            added += 1
    # 自动启动后台识别管线
    if added > 0:
        import time
        global _recognize_task
        _scan_progress["total"] = 0
        _scan_progress["pending"] = 0
        _scan_progress["done"] = 0
        _scan_progress["failed"] = 0
        _scan_progress["skipped"] = 0
        _scan_progress["current"] = None
        _scan_progress["status"] = "starting"
        _scan_progress["cancel_flag"] = False
        _scan_progress["pause_flag"] = False
        _scan_progress["cancel_categories"] = []
        _scan_progress["task_type"] = "recognize"
        _scan_progress["started_at"] = time.strftime("%Y-%m-%d %H:%M:%S")
        _scan_progress["completed_at"] = None
        _scan_progress["current_category"] = ""
        _scan_progress["done_categories"] = 0
        _scan_progress["categories"] = []
        asyncio.ensure_future(_safe_cancel_task()).add_done_callback(lambda _: _start_recognition())
    return {"status": "ok", "category": name, "added": added, "auto_recognize": added > 0}


# ===== 子分类管理 API (Phase 4) =====

@app.get("/api/image-index/subcategories")
async def api_get_subcategories():
    """获取所有子分类（含描述+图片数）"""
    conn = get_db()
    cats = conn.execute("""
        SELECT ic.category_name as name, 
               COUNT(ic.image_id) as count,
               cp.description,
               cp.created_at
        FROM image_categories ic
        LEFT JOIN category_profiles cp ON ic.category_name = cp.name
        LEFT JOIN image_index ii ON ic.image_id = ii.id
        WHERE ii.status != 'deleted'
        GROUP BY ic.category_name
        ORDER BY ic.category_name
    """).fetchall()
    conn.close()
    return [dict(c) for c in cats]


@app.put("/api/image-index/subcategories/{name}")
async def api_update_subcategory(name: str, data: dict = Body(...)):
    """更新子分类描述"""
    description = data.get("description", "")
    conn = get_db()
    conn.execute("""
        INSERT OR REPLACE INTO category_profiles (name, description, updated_at)
        VALUES (?, ?, datetime('now','localtime'))
    """, (name, description))
    conn.commit()
    conn.close()
    # 异步更新该分类的向量（非阻塞）
    try:
        from database import rebuild_category_profile_vectors
        import threading
        threading.Thread(target=rebuild_category_profile_vectors, daemon=True).start()
    except Exception:
        pass
    return {"status": "ok"}


@app.post("/api/image-index/rebuild-category-vectors")
async def api_rebuild_category_vectors():
    """批量重建所有分类描述的向量索引"""
    try:
        result = rebuild_category_profile_vectors()
        return {"status": "ok", "detail": result}
    except Exception as e:
        return {"status": "error", "detail": str(e)}


@app.delete("/api/image-index/subcategories/{name}")
async def api_delete_subcategory(name: str):
    """删除子分类+所有关联"""
    conn = get_db()
    # 获取所有关联的图片ID（通过 image_categories 或 image_index.category）
    cur = conn.execute("""
        SELECT DISTINCT ii.id FROM image_index ii
        LEFT JOIN image_categories ic ON ii.id = ic.image_id
        WHERE ic.category_name=? OR ii.category=?
    """, (name, name))
    deleted_ids = [row['id'] for row in cur.fetchall()]

    conn.execute("DELETE FROM image_categories WHERE category_name=?", (name,))
    conn.execute("DELETE FROM category_profiles WHERE name=?", (name,))
    # 删除所有关联图片
    if deleted_ids:
        placeholders = ','.join('?' * len(deleted_ids))
        conn.execute(f"DELETE FROM image_index WHERE id IN ({placeholders})", deleted_ids)
    conn.commit()
    conn.close()

    # 清理 image_vectors.db 中对应的向量
    if deleted_ids:
        try:
            from knowledge import _get_image_vec_conn
            vconn = _get_image_vec_conn()
            placeholders = ','.join('?' * len(deleted_ids))
            vconn.execute(f"DELETE FROM image_vectors WHERE image_id IN ({placeholders})", deleted_ids)
            vconn.commit()
            vconn.close()
        except Exception:
            pass
    return {"status": "ok", "deleted": True, "name": name, "image_count": len(deleted_ids)}


@app.get("/api/image-index/subcategories/{name}/images")
async def api_get_category_images_route(
    name: str,
    q: str = "",
    sort: str = "created_at",
    order: str = "desc",
    offset: int = 0,
    limit: int = 50
):
    """获取子分类内图片（支持搜索/排序/分页）"""
    images = get_category_images(name, limit=limit, offset=offset, q=q, sort=sort, order=order)
    conn = get_db()
    total = conn.execute(
            "SELECT COUNT(*) as c FROM image_categories ic "
            "JOIN image_index ii ON ic.image_id = ii.id "
            "WHERE ic.category_name=? AND ii.status!='deleted'",
            (name,)
        ).fetchone()["c"]
    conn.close()
    return {"images": images, "total": total}


@app.get("/api/image-index/subcategories/{name}/images/{image_id}/position")
async def api_get_image_position(name: str, image_id: int, sort: str = "created_at", order: str = "desc"):
    """获取图片在分类排序中的位置（第几条），用于前端跳转定位。
    直接查出排序后的 id 列表，找目标 id 的索引，保证和列表显示顺序完全一致。"""
    conn = get_db()
    
    # 排序规则必须和 get_category_images 完全一致
    if sort == "id_asc":
        sort_col = "ii.id"
        order_sql = "ASC"
    elif sort == "id_desc":
        sort_col = "ii.id"
        order_sql = "DESC"
    elif sort == "use_count":
        sort_col = "ii.use_count"
        order_sql = "DESC"
    else:  # created_at
        sort_col = "ii.created_at"
        order_sql = "DESC" if order == "desc" else "ASC"
    
    # 查出该分类排序后的所有 id
    rows = conn.execute(
        f"SELECT ii.id FROM image_index ii "
        f"JOIN image_categories ic ON ii.id = ic.image_id "
        f"WHERE ic.category_name=? AND ii.status!='deleted' "
        f"ORDER BY {sort_col} {order_sql}, ii.id",
        (name,)
    ).fetchall()
    
    # 找目标 id 的位置
    position = 0
    for i, row in enumerate(rows):
        if row["id"] == image_id:
            position = i
            break
    
    conn.close()
    return {"position": position}


# ===== 单张图片 CRUD =====

@app.put("/api/image-index/images/{image_id}")
async def api_update_image(image_id: int, data: dict = Body(...)):
    """更新图片字段（描述/标签/适合客户/案例组/备注）"""
    conn = get_db()
    updates = []
    params = []
    for field in ["description", "tags", "applicable_customers", "case_group_id"]:
            if field in data:
                val = data[field]
                # 列表字段转 JSON 字符串存储
                if isinstance(val, list):
                    val = json.dumps(val, ensure_ascii=False)
                updates.append(f"{field}=?")
                params.append(val)
    if not updates:
        conn.close()
        raise HTTPException(status_code=400, detail="无有效更新字段")
    updates.append("updated_at=datetime('now','localtime')")
    params.append(image_id)
    conn.execute(f"UPDATE image_index SET {', '.join(updates)} WHERE id=?", params)
    conn.commit()
    conn.close()
    return {"status": "ok"}


@app.delete("/api/image-index/images/{image_id}")
async def api_delete_image(image_id: int):
    """删除图片+关联"""
    conn = get_db()
    conn.execute("DELETE FROM image_categories WHERE image_id=?", (image_id,))
    row = conn.execute("SELECT file_path FROM image_index WHERE id=?", (image_id,)).fetchone()
    conn.execute("DELETE FROM image_index WHERE id=?", (image_id,))
    # 删除向量
    try:
        from database import _get_image_vec_conn
        vconn = _get_image_vec_conn()
        vconn.execute("DELETE FROM image_vectors WHERE image_id=?", (image_id,))
        vconn.commit()
        vconn.close()
    except:
        pass
    conn.commit()
    conn.close()
    return {"status": "ok", "deleted": True, "image_id": image_id}


# ===== 上传/复制 =====

@app.post("/api/image-index/upload")
async def api_upload_images(files: list[UploadFile] = File(...), category: str = Form(...)):
    """上传本地图片到指定分类->识别->索引->向量"""
    import shutil
    from pathlib import Path
    from datetime import datetime

    IMAGES_DIR = BASE_DIR / "data" / "uploaded_images"
    IMAGES_DIR.mkdir(parents=True, exist_ok=True)

    results = []
    for file in files:
        if not file.filename:
            continue
        ext = Path(file.filename).suffix.lower()
        save_name = f"{datetime.now().strftime('%Y%m%d%H%M%S')}_{len(results)}{ext}"
        save_path = IMAGES_DIR / save_name
        content = await file.read()
        with open(save_path, "wb") as f:
            f.write(content)

        conn = get_db()
        cursor = conn.execute(
            "INSERT INTO image_index (file_path, category, status, file_size, created_at) "
            "VALUES (?, ?, 'pending', ?, datetime('now','localtime'))",
            (str(save_path), category, os.path.getsize(save_path))
        )
        img_id = cursor.lastrowid
        conn.execute("INSERT OR IGNORE INTO image_categories (image_id, category_name) VALUES (?, ?)",
                     (img_id, category))
        conn.commit()
        conn.close()
        results.append({"id": img_id, "file_path": str(save_path), "category": category})

    # 自动触发批量识别
    if results:
        global _recognize_task
        _scan_progress["total"] = 0
        _scan_progress["pending"] = 0
        _scan_progress["done"] = 0
        _scan_progress["failed"] = 0
        _scan_progress["skipped"] = 0
        _scan_progress["status"] = "starting"
        _scan_progress["cancel_flag"] = False
        _scan_progress["pause_flag"] = False
        _scan_progress["task_type"] = "recognize"
        asyncio.ensure_future(_safe_cancel_task()).add_done_callback(lambda _: _start_recognition())

    return {"status": "ok", "images": results}


@app.post("/api/image-index/copy-to-category")
async def api_copy_to_category(data: dict = Body(...)):
    """从其他分类复制图片到指定分类"""
    image_ids = data.get("image_ids", [])
    target_category = data.get("category", "")
    if not image_ids or not target_category:
        raise HTTPException(status_code=400, detail="请提供image_ids和category")
    conn = get_db()
    for img_id in image_ids:
        conn.execute("INSERT OR IGNORE INTO image_categories (image_id, category_name) VALUES (?, ?)",
                     (img_id, target_category))
    conn.commit()
    conn.close()
    return {"status": "ok", "copied": len(image_ids)}


# ===== 批量操作 =====

@app.post("/api/image-index/batch-delete")
async def api_batch_delete(data: dict = Body(...)):
    """批量删除图片"""
    image_ids = data.get("image_ids", [])
    if not image_ids:
        raise HTTPException(status_code=400, detail="请提供image_ids")
    conn = get_db()
    for img_id in image_ids:
        conn.execute("DELETE FROM image_categories WHERE image_id=?", (img_id,))
        conn.execute("DELETE FROM image_index WHERE id=?", (img_id,))
    conn.commit()
    conn.close()
    return {"status": "ok", "deleted": len(image_ids)}


@app.post("/api/image-index/batch-tag")
async def api_batch_tag(data: dict = Body(...)):
    """批量添加标签"""
    image_ids = data.get("image_ids", [])
    tags = data.get("tags", "")
    if not image_ids:
        raise HTTPException(status_code=400, detail="请提供image_ids")
    conn = get_db()
    for img_id in image_ids:
        row = conn.execute("SELECT tags FROM image_index WHERE id=?", (img_id,)).fetchone()
        if row:
            current = row["tags"]
            try:
                current_list = json.loads(current) if current and current != "[]" else []
            except:
                current_list = []
            new_tags_list = [t.strip() for t in tags.split(",") if t.strip()]
            merged = list(set(current_list + new_tags_list))
            conn.execute("UPDATE image_index SET tags=?, updated_at=datetime('now','localtime') WHERE id=?",
                         (json.dumps(merged, ensure_ascii=False), img_id))
    conn.commit()
    conn.close()
    return {"status": "ok", "updated": len(image_ids)}


# ===== 话术图片匹配 =====

@app.post("/api/image-index/match-for-script")
async def api_match_images_for_script(data: dict = Body(...)):
    """批量搜索图片索引匹配话术中的[生成图片:描述]"""
    descriptions = data.get("descriptions", [])
    results = {}
    for desc in descriptions:
        matches = search_images_hybrid(query=desc, limit=1)
        if matches:
            results[desc] = {"found": True, "image": matches[0]}
        else:
            results[desc] = {"found": False}
    return {"matches": results}


@app.get("/api/image-index/match")
async def api_image_index_match(needs: str = Query(...)):
    """匹配话术图片"""
    need_list = [n.strip() for n in needs.split(",") if n.strip()]
    return {"status": "ok", "data": match_images_for_script(need_list)}

@app.get("/api/image-index/{image_id}")
async def api_image_index_detail(image_id: int):
    """获取图片详情"""
    detail = get_image_detail(image_id)
    if not detail:
        raise HTTPException(status_code=404, detail="图片不存在")
    return {"status": "ok", "data": detail}

@app.get("/api/image-index/case/{case_group_id}")
async def api_image_index_case(case_group_id: str):
    """获取案例分组图片"""
    return {"status": "ok", "data": get_case_group(case_group_id)}

@app.get("/api/image-index/images/{image_id}")
async def api_image_index_serve_image(image_id: int):
    """提供图片文件"""
    detail = get_image_detail(image_id)
    if not detail:
        raise HTTPException(status_code=404, detail="图片不存在")
    file_path = detail.get("file_path", "")
    if not file_path:
        raise HTTPException(status_code=404, detail="文件不存在")
    # 支持相对路径（相对于项目根目录）
    if not os.path.isabs(file_path):
        file_path = str(BASE_DIR / file_path)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="文件不存在")
    ext = Path(file_path).suffix.lower()
    mime_map = {'.jpg': 'image/jpeg', '.jpeg': 'image/jpeg', '.png': 'image/png', '.gif': 'image/gif', '.webp': 'image/webp'}
    return FileResponse(file_path, media_type=mime_map.get(ext, 'application/octet-stream'))


# ===== 案例归组 API =====

@app.post("/api/image-index/group/auto")
async def api_image_index_group_auto():
    """自动归组（时间戳连续<30秒）"""
    groups = auto_group_cases()
    valid_groups = {k: [m['id'] for m in v] for k, v in groups.items() if v}
    return {"status": "ok", "data": {"groups": valid_groups, "count": len(valid_groups)}}


@app.post("/api/image-index/group/confirm/{group_id}")
async def api_image_index_group_confirm(group_id: str):
    """确认归组（更新数据库）"""
    groups = auto_group_cases()
    if group_id not in groups:
        raise HTTPException(status_code=404, detail="组不存在")
    image_ids = [m['id'] for m in groups[group_id]]
    assign_case_group(group_id, image_ids)
    return {"status": "ok", "data": {"group_id": group_id, "image_count": len(image_ids)}}


@app.post("/api/image-index/group/llm/{group_id}")
async def api_image_index_group_llm(group_id: str):
    """LLM二次确认归组"""
    groups = auto_group_cases()
    if group_id not in groups:
        raise HTTPException(status_code=404, detail="组不存在")
    image_ids = [m['id'] for m in groups[group_id]]
    details = [get_image_detail(iid) for iid in image_ids]
    details = [d for d in details if d]
    from llm_client import _confirm_case_group
    result = await _confirm_case_group(details, group_id)
    return {"status": "ok", "data": result}


@app.get("/api/image-index/ungrouped")
async def api_image_index_ungrouped(limit: int = 50):
    """获取未归组的图片"""
    return {"status": "ok", "data": get_ungrouped_images(limit)}


@app.post("/api/image-index/group/assign")
async def api_image_index_group_assign(data: dict = Body(...)):
    """手动分配案例组"""
    group_id = data.get("group_id", "")
    image_ids = data.get("image_ids", [])
    if not group_id or not image_ids:
        raise HTTPException(status_code=400, detail="缺少group_id或image_ids")
    assign_case_group(group_id, image_ids)
    return {"status": "ok", "data": {"group_id": group_id, "image_count": len(image_ids)}}







# ===== 首页 =====

FRONTEND_DIR = (Path(__file__).resolve().parent.parent / "frontend").resolve()

# 提供前端图标文件
@app.get("/logo.jpeg")
async def serve_logo():
    """提供logo图片"""
    logo_path = FRONTEND_DIR / "logo.jpeg"
    if logo_path.exists():
        return FileResponse(str(logo_path), media_type="image/jpeg")
    return HTMLResponse(content="", status_code=204)

@app.get("/favicon.ico")
async def serve_favicon():
    """提供favicon"""
    icon_path = FRONTEND_DIR / "favicon.svg"
    if icon_path.exists():
        return FileResponse(str(icon_path), media_type="image/svg+xml")
    return HTMLResponse(content="", status_code=204)

@app.get("/frontend/icons/{filename}")
async def serve_icon(filename: str):
    """提供前端图标文件"""
    icon_path = FRONTEND_DIR / "icons" / filename
    if icon_path.exists():
        return FileResponse(str(icon_path))
    return HTMLResponse(content="", status_code=404)

@app.get("/")
async def serve_frontend():
    """返回前端首页"""
    index_path = FRONTEND_DIR / "index.html"
    if index_path.exists():
        return HTMLResponse(content=index_path.read_text(encoding="utf-8"))
    return HTMLResponse("<h1>小赛助手</h1><p>前端文件未找到</p>")



# ===== 启动入口 =====

def get_local_ip():
    """获取本机局域网IP"""
    import socket
    try:
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        s.connect(("8.8.8.8", 80))
        ip = s.getsockname()[0]
        s.close()
        return ip
    except Exception:
        return "127.0.0.1"


def start_server():
    """启动服务（自动找空闲端口）"""
    config = cfg.load_config()
    base_port = config["app"]["port"]

    # 从配置端口开始试，找到空闲端口
    port = base_port
    while True:
        try:
            import socket
            s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
            s.settimeout(0.5)
            result = s.connect_ex(('127.0.0.1', port))
            s.close()
            if result != 0:
                break  # 端口空闲
            port += 1
        except:
            port += 1

    if port != base_port:
        print(f"⚠️ 端口 {base_port} 被占用，自动切换到端口 {port}")

    print(f"[小赛助手 v2] 启动中...")
    print(f"   访问地址: http://localhost:{port}")
    print(f"   局域网访问: http://{get_local_ip()}:{port}")
    print(f"   按 Ctrl+C 停止服务")
    print()

    # 先启动服务器，等就绪后再打开浏览器
    # 先启动服务器，等就绪后再打开浏览器
    import threading
    def _open_browser_when_ready():
        import time, urllib.request
        for _ in range(30):
            time.sleep(1)
            try:
                urllib.request.urlopen(f"http://localhost:{port}", timeout=2)
                break
            except:
                continue
        try:
            import subprocess, webbrowser
            chrome_paths = [
                r"C:\Program Files\Google\Chrome\Application\chrome.exe",
                r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
                os.path.expanduser(r"~\AppData\Local\Google\Chrome\Application\chrome.exe"),
            ]
            chrome_found = None
            for p in chrome_paths:
                if os.path.exists(p):
                    chrome_found = p
                    break
            if chrome_found:
                subprocess.Popen([chrome_found, f"http://localhost:{port}"])
            else:
                webbrowser.open(f"http://localhost:{port}")
        except Exception:
            webbrowser.open(f"http://localhost:{port}")
    threading.Thread(target=_open_browser_when_ready, daemon=True).start()

    uvicorn.run(app, host="0.0.0.0", port=port, log_level="info")

# ===== 图片索引分类 API（新增） =====

@app.get("/api/image-index/browse/{category_name}")
async def api_image_index_category(category_name: str, limit: int = 500):
    """获取指定分类的所有图片（含描述、标签、适用客户）"""
    from database import get_category_images
    return {"status": "ok", "data": get_category_images(category_name, limit)}


@app.get("/api/image-index/config")
async def api_image_index_config():
    """获取所有分类的配置信息（供前端展示）"""
    from prompts import CATEGORY_CONFIG, get_all_categories
    categories = get_all_categories()
    configs = {}
    for cat in categories:
        cfg = CATEGORY_CONFIG.get(cat, {})
        configs[cat] = {
            "type": cfg.get("type", "通用"),
            "tag_keys": cfg.get("tag_keys", []),
            "scenarios": [s[1] for s in cfg.get("default_scenarios", [])],
        }
    return {"status": "ok", "data": {"categories": categories, "configs": configs}}


@app.post("/api/image-index/re-recognize/{image_id}")
async def api_image_index_rerecognize(image_id: int):
    """重新识别单张图片（轮询模型）"""
    from database import get_db
    conn = get_db()
    row = conn.execute("SELECT * FROM image_index WHERE id=?", (image_id,)).fetchone()
    conn.close()
    if not row:
        raise HTTPException(status_code=404, detail="图片不存在")
    img = dict(row)
    from llm_client import _analyze_image_with_model
    from config_manager import get_vision_model_list
    models = get_vision_model_list()
    if not models:
        from config_manager import get_vision_model_config
        models = [get_vision_model_config()]
    last_error = None
    for m in models:
        try:
            desc, tags, customers = await _analyze_image_with_model(
                img["file_path"], img.get("category", ""), m)
            from database import update_image_result
            update_image_result(image_id, desc, tags, applicable_customers=customers)
            return {"status": "ok", "data": {"description": desc, "tags": tags, "applicable_customers": customers}}
        except Exception as e:
            last_error = e
            continue
    raise HTTPException(status_code=500, detail=f"所有模型均失败: {last_error}")


@app.post("/api/image-index/re-recognize-failed")
async def api_image_index_rerecognize_failed():
    """重新识别所有失败图片（改为pending，触发批量识别）"""
    from database import get_db
    conn = get_db()
    failed = conn.execute("SELECT COUNT(*) FROM image_index WHERE status='failed'").fetchone()[0]
    if failed == 0:
        conn.close()
        return {"status": "ok", "message": "暂无待识别图片", "reset": 0}
    conn.execute("UPDATE image_index SET status='pending', error_msg='', updated_at=datetime('now','localtime') WHERE status='failed'")
    conn.commit()
    conn.close()
    global _recognize_task
    _scan_progress["total"] = 0
    _scan_progress["pending"] = 0
    _scan_progress["done"] = 0
    _scan_progress["failed"] = 0
    _scan_progress["skipped"] = 0
    _scan_progress["status"] = "starting"
    _scan_progress["cancel_flag"] = False
    _scan_progress["pause_flag"] = False
    _scan_progress["task_type"] = "recognize"
    asyncio.ensure_future(_safe_cancel_task()).add_done_callback(lambda _: _start_recognition())
    return {"status": "ok", "message": f"已重置 {failed} 张失败图片", "reset": failed}


@app.post("/api/image-index/task/delete-images")
async def api_image_index_delete_images(data: dict = Body(...)):
    """删除指定ID列表的图片（用于取消任务时清理）"""
    ids = data.get("ids", [])
    if not ids:
        return {"status": "ok", "deleted": 0}
    placeholders = ','.join('?' * len(ids))
    conn = get_db()
    conn.execute(f"DELETE FROM image_categories WHERE image_id IN ({placeholders})", ids)
    conn.execute(f"DELETE FROM image_index WHERE id IN ({placeholders})", ids)
    conn.commit()
    conn.close()
    try:
        from knowledge import _get_image_vec_conn
        vconn = _get_image_vec_conn()
        vconn.execute(f"DELETE FROM image_vectors WHERE image_id IN ({placeholders})", ids)
        vconn.commit()
        vconn.close()
    except:
        pass
    return {"status": "ok", "deleted": len(ids)}


@app.post("/api/image-index/failed/ignore")
async def api_image_index_ignore_failed():
    """删除所有失败图片"""
    conn = get_db()
    ids = [r["id"] for r in conn.execute("SELECT id FROM image_index WHERE status='failed'").fetchall()]
    if not ids:
        conn.close()
        return {"status": "ok", "deleted": 0}
    placeholders = ','.join('?' * len(ids))
    conn.execute(f"DELETE FROM image_categories WHERE image_id IN ({placeholders})", ids)
    conn.execute(f"DELETE FROM image_index WHERE id IN ({placeholders})", ids)
    conn.commit()
    conn.close()
    try:
        from knowledge import _get_image_vec_conn
        vconn = _get_image_vec_conn()
        vconn.execute(f"DELETE FROM image_vectors WHERE image_id IN ({placeholders})", ids)
        vconn.commit()
        vconn.close()
    except:
        pass
    return {"status": "ok", "deleted": len(ids)}


@app.post("/api/image-index/clear-all")
async def api_image_index_clear_all():
    """清空所有图片索引数据"""
    from database import get_db, DB_IMAGE_VEC_PATH, DB_VECTORS_PATH
    conn = get_db()
    conn.execute("DELETE FROM image_index")
    conn.execute("DELETE FROM image_categories")
    conn.execute("DELETE FROM category_profiles")
    conn.commit()
    conn.close()
    # 清空向量数据库
    for vec_db in [DB_IMAGE_VEC_PATH, DB_VECTORS_PATH]:
        try:
            import sqlite3
            vconn = sqlite3.connect(str(vec_db))
            for table in ['image_vectors', 'category_profile_vectors', 'knowledge_vectors']:
                try:
                    vconn.execute(f"DELETE FROM {table}")
                except:
                    pass
            vconn.commit()
            vconn.close()
        except:
            pass
    return {"status": "ok", "message": "已清空所有图片数据"}


@app.post("/api/image-index/re-scan")
async def api_image_index_rescan(data: dict = Body(...)):
    """重新扫描某个分类（新增分类时使用）"""
    folder_path = data.get("folder_path", "")
    category_name = data.get("category_name", "")
    if not folder_path or not category_name:
        raise HTTPException(status_code=400, detail="需要folder_path和category_name")
    from pathlib import Path
    from database import get_db
    results = []
    p = Path(folder_path)
    if p.is_dir():
        for f in sorted(p.iterdir()):
            if f.suffix.lower() in ('.jpg','.jpeg','.png','.gif','.webp','.bmp'):
                fp = str(f)
                try:
                    import hashlib
                    with open(fp, 'rb') as fh:
                        fhash = hashlib.md5(fh.read()).hexdigest()
                except:
                    fhash = ""
                fsize = f.stat().st_size
                results.append({
                    "file_path": fp, "file_hash": fhash,
                    "category": category_name, "status": "pending",
                    "file_size": fsize, "filename": f.name
                })
    if results:
        conn = get_db()
        for r in results:
            conn.execute(
                "INSERT INTO image_index (file_path, file_hash, category, status, file_size) VALUES (?, ?, ?, 'pending', ?)",
                (r["file_path"], r["file_hash"], r["category"], r["file_size"])
            )
        conn.commit()
        conn.close()
    return {"status": "ok", "data": {"category": category_name, "image_count": len(results)}}



# ===== 话术模板 API =====

@app.get("/api/script-templates")
async def api_list_script_templates(search: str = Query(""), category: str = Query("")):
    """获取话术模板列表"""
    from database import list_script_templates
    templates = list_script_templates(search=search, category=category)
    return {"templates": templates}


@app.post("/api/script-templates")
async def api_create_script_template(data: dict = Body(...)):
    """创建话术模板"""
    from database import add_script_template
    title = data.get("title", "").strip()
    content = data.get("content", "").strip()
    if not title or not content:
        raise HTTPException(status_code=400, detail="标题和内容不能为空")
    tags = data.get("tags", [])
    category = data.get("category", "")
    template_id = add_script_template(title, content, tags, category)
    return {"status": "ok", "id": template_id}


@app.put("/api/script-templates/{template_id}")
async def api_update_script_template(template_id: int, data: dict = Body(...)):
    """更新话术模板"""
    from database import update_script_template
    update_script_template(
        template_id,
        title=data.get("title"),
        content=data.get("content"),
        tags=data.get("tags"),
        category=data.get("category")
    )
    return {"status": "ok"}


@app.delete("/api/script-templates/{template_id}")
async def api_delete_script_template(template_id: int):
    """删除话术模板"""
    from database import delete_script_template
    delete_script_template(template_id)
    return {"status": "ok"}


@app.get("/api/script-templates/{template_id}")
async def api_get_script_template(template_id: int):
    """获取单个话术模板"""
    from database import get_script_template
    template = get_script_template(template_id)
    if not template:
        raise HTTPException(status_code=404, detail="模板不存在")
    return {"template": template}


# ===== 话术导出Word =====

@app.get("/api/script-templates/{template_id}/export")
async def api_export_script_template_word(template_id: int):
    """导出话术模板为Word文档"""
    from database import get_script_template
    template = get_script_template(template_id)
    if not template:
        raise HTTPException(status_code=404, detail="模板不存在")
    
    try:
        import docx
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # 标题
        title = doc.add_heading(template["title"], level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 标签
        tags = template.get("tags", [])
        if isinstance(tags, str):
            import json
            tags = json.loads(tags) if tags else []
        if tags:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run("标签: " + "、".join(tags))
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(100, 100, 100)
        
        # 分隔线
        doc.add_paragraph("─" * 40)
        
        # 内容
        content = template["content"]
        # 按行解析，支持分段
        lines = content.split("\n")
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph("")  # 空行
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=2)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=3)
            elif line.startswith("---"):
                doc.add_paragraph("─" * 40)
            elif line.startswith("- "):
                doc.add_paragraph(line, style="List Bullet")
            elif "[生成图片:" in line:
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.italic = True
                run.font.color.rgb = RGBColor(0, 100, 200)
            else:
                doc.add_paragraph(line)
        
        # 页脚信息
        doc.add_paragraph("")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"— 由小赛助手生成 —")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(150, 150, 150)
        
        # 返回文件
        import io
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        
        safe_title = "".join(c for c in template["title"] if c.isalnum() or c in " \u4e00-\u9fff\u3000-\u303f")
        filename = f"{safe_title}.docx"
        
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )
    except ImportError:
        raise HTTPException(status_code=500, detail="导出Word需要安装python-docx: pip install python-docx")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"导出失败: {str(e)}")


# ===== 有效话术库 =====

@app.get("/api/effective-scripts")
async def api_list_effective_scripts(scenario: str = Query(""), customer_type: str = Query(""),
                                      sort_by: str = Query("effective_count"), sort_order: str = Query("desc"),
                                      limit: int = Query(100), offset: int = Query(0)):
    """获取有效话术列表"""
    return list_effective_scripts(scenario, customer_type, sort_by, sort_order, limit, offset)


@app.get("/api/effective-scripts/stats")
async def api_effective_script_stats():
    """获取有效话术统计"""
    return get_effective_script_stats()


@app.delete("/api/effective-scripts/{script_id}")
async def api_delete_effective_script(script_id: int):
    """删除有效话术"""
    if delete_effective_script(script_id):
        return {"status": "ok"}
    raise HTTPException(status_code=404, detail="话术不存在")


@app.post("/api/effective-scripts/dedup")
async def api_dedup_scripts():
    """去重合并"""
    result = dedup_effective_scripts()
    return {"status": "ok", "result": result}


@app.put("/api/effective-scripts/{script_id}")
async def api_update_script(script_id: int, data: dict = Body(...)):
    """编辑有效话术"""
    update_effective_script(script_id, data.get("content", ""), data.get("scenario", ""), data.get("customer_type", ""))
    return {"status": "ok"}


# ===== 常见问题 =====

@app.get("/api/faq")
async def api_list_faqs(category: str = Query("")):
    """获取FAQ列表"""
    return {"faqs": list_faqs(category)}


@app.get("/api/faq/{faq_id}")
async def api_get_faq(faq_id: int):
    """获取单个FAQ"""
    faq = get_faq(faq_id)
    if not faq:
        raise HTTPException(status_code=404, detail="FAQ不存在")
    return {"faq": faq}


@app.post("/api/faq")
async def api_add_faq(data: dict = Body(...)):
    """添加FAQ"""
    question = data.get("question", "").strip()
    answer = data.get("answer", "").strip()
    if not question or not answer:
        raise HTTPException(status_code=400, detail="问题和答案不能为空")
    result = add_faq(question, answer, data.get("category", ""), data.get("sort_order", 0))
    return {"status": "ok", "faq": result}


@app.put("/api/faq/{faq_id}")
async def api_update_faq(faq_id: int, data: dict = Body(...)):
    """更新FAQ"""
    faq = get_faq(faq_id)
    if not faq:
        raise HTTPException(status_code=404, detail="FAQ不存在")
    update_faq(faq_id, data.get("question", ""), data.get("answer", ""),
               data.get("category", ""), data.get("sort_order", 0))
    return {"status": "ok"}


@app.delete("/api/faq/{faq_id}")
async def api_delete_faq(faq_id: int):
    """删除FAQ"""
    if delete_faq(faq_id):
        return {"status": "ok"}
    raise HTTPException(status_code=404, detail="FAQ不存在")


# ===== 数据统计 API =====

@app.get("/api/stats")
async def api_stats():
    """获取数据统计（跨库：主库/messages库/向量库）"""
    from database import get_db, get_messages_db, get_vectors_db
    from datetime import datetime, timedelta

    # 主库：customers / image_index
    main_conn = get_db()
    # 消息库：messages
    msg_conn = get_messages_db()
    # 向量库：knowledge_chunks（文本元数据实际存于向量库）
    vec_conn = None
    try:
        vec_conn = get_vectors_db()
    except Exception:
        pass

    try:
        # 客户总数
        customer_count = main_conn.execute("SELECT COUNT(*) as cnt FROM customers").fetchone()["cnt"]

        # 消息总数
        message_count = msg_conn.execute("SELECT COUNT(*) as cnt FROM messages").fetchone()["cnt"]

        # 今日消息数
        today = datetime.now().strftime("%Y-%m-%d")
        today_msg_count = msg_conn.execute("SELECT COUNT(*) as cnt FROM messages WHERE timestamp LIKE ?", (f"{today}%",)).fetchone()["cnt"]

        # 知识库文档数 / chunks 数
        doc_count = 0
        chunk_count = 0
        if vec_conn is not None:
            try:
                doc_count = vec_conn.execute("SELECT COUNT(DISTINCT doc_id) as cnt FROM knowledge_chunks").fetchone()["cnt"]
                chunk_count = vec_conn.execute("SELECT COUNT(*) as cnt FROM knowledge_chunks").fetchone()["cnt"]
            except Exception:
                pass

        # 图片索引数
        image_count = main_conn.execute("SELECT COUNT(*) as cnt FROM image_index").fetchone()["cnt"]

        # 客户活跃度（最近7天有消息的客户数）
        week_ago = (datetime.now() - timedelta(days=7)).strftime("%Y-%m-%d %H:%M:%S")
        active_customers = msg_conn.execute("SELECT COUNT(DISTINCT customer_id) as cnt FROM messages WHERE timestamp > ?", (week_ago,)).fetchone()["cnt"]

        return {
            "customer_count": customer_count,
            "message_count": message_count,
            "today_msg_count": today_msg_count,
            "doc_count": doc_count,
            "chunk_count": chunk_count,
            "image_count": image_count,
            "active_customers_7d": active_customers,
            "stats_time": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        }
    finally:
        main_conn.close()
        msg_conn.close()
        if vec_conn is not None:
            vec_conn.close()

if __name__ == "__main__":
    start_server()


