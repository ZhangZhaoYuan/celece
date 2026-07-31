"""
知识库管理 - sqlite-vec 向量检索（语义搜索）
v2.0 改用 sqlite-vec 替代 ChromaDB，使用统一个 SQLite 文件存储向量
向量模型：DashScope text-embedding-v2（API 调用，无需本地模型）

混合检索（Hybrid Search）：语义搜索 + 关键词搜索，RRF 加权融合
当 embedding API 不可用时，自动降级为纯关键词搜索
"""
import os
import json
import hashlib
import shutil
import sqlite3
import time
import urllib.request
from pathlib import Path
from datetime import datetime
from typing import Optional, List

BASE_DIR = Path(__file__).resolve().parent.parent
DATA_DIR = BASE_DIR / "data"
KNOWLEDGE_DIR = DATA_DIR / "knowledge"
DB_PATH = DATA_DIR / "customers.db"  # 主库
VEC_DB_PATH = DATA_DIR / "vectors.db"  # 向量库（独立文件）

# 全局标记：embedding API 是否可用
_embedding_available = True

# 从配置读取 embedding 模型参数
def _init_embedding_config():
    """从 config.json 加载 embedding 配置"""
    global _embedding_model, _embedding_dim, _embedding_url, _embedding_api_mode
    try:
        import config_manager as cfg
        emb = cfg.get_embedding_config()
        _embedding_model = emb.get("model", "text-embedding-v2")
        _embedding_dim = emb.get("dim", 1536)
        _embedding_url = emb.get("base_url", "https://dashscope.aliyuncs.com/compatible-mode/v1/embeddings")
        _embedding_api_mode = emb.get("api_mode", "openai")
    except Exception:
        _embedding_model = "text-embedding-v2"
        _embedding_dim = 1536
        _embedding_url = "https://dashscope.aliyuncs.com/compatible-mode/v1/embeddings"
        _embedding_api_mode = "openai"

_embedding_model = "text-embedding-v2"
_embedding_dim = 1536
_embedding_url = "https://dashscope.aliyuncs.com/compatible-mode/v1/embeddings"
_embedding_api_mode = "openai"
_init_embedding_config()

# 错误重试机制：瞬态错误 5 分钟后自动复位重试
_embedding_failure_time = 0  # 上次失败的时间戳
_EMBEDDING_RETRY_INTERVAL = 300  # 5 分钟（秒）

# ===== API Key 获取 =====

def _get_api_key() -> str:
    try:
        import config_manager as cfg
        # 优先从 embedding 模型配置读取 key
        emb = cfg.get_embedding_config()
        key = emb.get("api_key", "")
        if key:
            return key
        # fallback
        config = cfg.load_config()
        key = config.get("embedding_api_key", "")
        if key:
            return key
        return config.get("llm", {}).get("api_key", "")
    except Exception:
        return ""


# ===== sqlite-vec 初始化 =====

def _get_vec_conn() -> sqlite3.Connection:
    """获取启用了 sqlite-vec 扩展的数据库连接（使用独立向量库文件）"""
    global _embedding_available
    conn = sqlite3.connect(str(VEC_DB_PATH))
    conn.enable_load_extension(True)
    try:
        # 尝试用 sqlite_vec.load() 加载（开发模式）
        import sqlite_vec
        sqlite_vec.load(conn)
        conn.enable_load_extension(False)
        conn.row_factory = sqlite3.Row
        return conn
    except (sqlite3.OperationalError, FileNotFoundError) as e:
        # PyInstaller 打包后 __file__ 指向 PYZ 虚拟路径，loadable_path() 找不到真实文件
        # 尝试手动定位 vec0.dll
        try:
            vec_dll = _find_vec_dll()
            if vec_dll and os.path.exists(vec_dll):
                conn.load_extension(vec_dll)
                conn.enable_load_extension(False)
                conn.row_factory = sqlite3.Row
                return conn
        except Exception:
            pass
        # 彻底降级
        _embedding_available = False
        conn.enable_load_extension(False)
        conn.row_factory = sqlite3.Row
        print(f"[sqlite-vec] 加载失败，降级为关键词搜索: {e}")
        return conn


def _find_vec_dll() -> str:
    """寻找 vec0.dll 的路径（处理 PyInstaller 打包后的路径偏移）"""
    import sys
    # 冻结模式：sys._MEIPASS 指向 _internal/
    if getattr(sys, 'frozen', False):
        candidates = [
            os.path.join(sys._MEIPASS, 'sqlite_vec', 'vec0.dll'),
            os.path.join(os.path.dirname(sys.executable), '_internal', 'sqlite_vec', 'vec0.dll'),
        ]
    else:
        candidates = []
    # 开发模式：用 sqlite_vec.__file__ 定位
    try:
        import sqlite_vec
        candidates.append(os.path.join(os.path.dirname(sqlite_vec.__file__), 'vec0.dll'))
    except ImportError:
        pass
    for c in candidates:
        if os.path.exists(c):
            return c
    return ""


def _ensure_db():
    """确保数据库表和向量索引已创建"""
    # knowledge_chunks 在主库
    mconn = sqlite3.connect(str(DB_PATH))
    mconn.row_factory = sqlite3.Row
    mconn.execute("""
        CREATE TABLE IF NOT EXISTS knowledge_chunks (
            id TEXT PRIMARY KEY,
            doc_id TEXT NOT NULL,
            chunk_index INTEGER NOT NULL,
            content TEXT NOT NULL,
            filename TEXT DEFAULT '',
            title TEXT DEFAULT '',
            created_at TEXT DEFAULT (datetime('now','localtime'))
        )
    """)
    mconn.commit()
    mconn.close()

    # sqlite-vec 向量索引表在向量库
    conn = _get_vec_conn()
    try:
        conn.execute(f"""
            CREATE VIRTUAL TABLE IF NOT EXISTS knowledge_vectors USING vec0(
                id TEXT PRIMARY KEY,
                vector FLOAT[{_embedding_dim}]
            )
        """)
    except Exception:
        pass  # 已存在则跳过

    conn.commit()
    conn.close()


# ===== Embedding API 调用 =====

def _get_embedding(text: str) -> Optional[List[float]]:
    """调用 embedding API 获取单条文本的向量"""
    global _embedding_available, _embedding_failure_time

    api_key = _get_api_key()
    if not api_key:
        _embedding_available = False
        _embedding_failure_time = time.time()
        return None

    try:
        # 截断到 2000 字符确保安全
        if len(text) > 2000:
            text = text[:2000]

        # 根据 api_mode 选择请求格式
        if _embedding_api_mode == "dashscope_multimodal":
            data = json.dumps({
                "model": _embedding_model,
                "input": {"contents": [{"text": text}]},
                "parameters": {"dimension": _embedding_dim}
            }).encode("utf-8")
        elif _embedding_api_mode == "dashscope_text":
            data = json.dumps({
                "model": _embedding_model,
                "input": {"texts": [text]}
            }).encode("utf-8")
        else:  # openai (默认)
            data = json.dumps({"input": text, "model": _embedding_model}).encode("utf-8")

        req = urllib.request.Request(
            _embedding_url, data=data,
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json"
            },
            method="POST"
        )
        resp = urllib.request.urlopen(req, timeout=15)
        result = json.loads(resp.read().decode("utf-8"))
        vector = _parse_embedding_response(result, _embedding_api_mode)

        # L2 归一化（使余弦相似度等价于 L2 距离）
        norm = sum(v * v for v in vector) ** 0.5
        if norm > 0:
            vector = [v / norm for v in vector]

        _embedding_available = True
        return vector

    except Exception as e:
        print(f"[知识库] embedding API 调用失败: {e}")
        try:
            body_text = json.dumps(data.decode("utf-8") if isinstance(data, bytes) else data)[:200]
        except:
            body_text = str(data)[:200]
        if hasattr(e, 'read'):
            try:
                resp_body = e.read().decode("utf-8", errors="replace")[:300]
                print(f"  [调试] 请求体: {body_text}")
                print(f"  [调试] 响应体: {resp_body}")
            except:
                pass
        _embedding_available = False
        _embedding_failure_time = time.time()
        return None


def _get_embeddings_batch(texts: List[str]) -> Optional[List[List[float]]]:
    """批量获取向量（自动按 10 条一批分片调用）"""
    global _embedding_available, _embedding_failure_time

    api_key = _get_api_key()
    if not api_key:
        _embedding_available = False
        _embedding_failure_time = time.time()
        return None

    # 截断每条文本到 2000 字符
    texts = [t[:2000] if len(t) > 2000 else t for t in texts]

    # 部分模型（如 text-embedding-v3）限制单次最多 10 条，自动分片
    BATCH_LIMIT = 10
    all_vectors = []

    for start in range(0, len(texts), BATCH_LIMIT):
        batch = texts[start:start + BATCH_LIMIT]
        try:
            data = json.dumps(_build_embedding_data_batch(batch, _embedding_model, _embedding_api_mode, _embedding_dim)).encode("utf-8")
            req = urllib.request.Request(
                _embedding_url, data=data,
                headers={
                    "Authorization": f"Bearer {api_key}",
                    "Content-Type": "application/json"
                },
                method="POST"
            )
            resp = urllib.request.urlopen(req, timeout=60)
            result = json.loads(resp.read().decode("utf-8"))

            for vec in _parse_embedding_response_batch(result, _embedding_api_mode):
                norm = sum(v * v for v in vec) ** 0.5
                if norm > 0:
                    vec = [v / norm for v in vec]
                all_vectors.append(vec)

        except Exception as e:
            print(f"[知识库] embedding API 批量调用失败 (批次 {start//BATCH_LIMIT}): {e}")
            _embedding_available = False
            _embedding_failure_time = time.time()
            return None

    _embedding_available = True
    return all_vectors


def is_embedding_available() -> bool:
    """外部查询 embedding API 状态（含超时自动复位）"""
    global _embedding_available
    if not _embedding_available and time.time() - _embedding_failure_time > _EMBEDDING_RETRY_INTERVAL:
        _embedding_available = True
    return _embedding_available


def _vec_to_str(vector: List[float]) -> str:
    """向量转 sqlite-vec 兼容的 JSON 字符串"""
    return "[" + ",".join(f"{v:.8f}" for v in vector) + "]"


def _build_embedding_data(text: str, model: str, api_mode: str, dim: int) -> dict:
    """根据 api_mode 构建 embedding API 请求体"""
    if api_mode == "dashscope_multimodal":
        return {
            "model": model,
            "input": {"contents": [{"text": text}]},
            "parameters": {"dimension": dim}
        }
    elif api_mode == "dashscope_text":
        return {
            "model": model,
            "input": {"texts": [text]}
        }
    else:  # openai
        return {"input": text, "model": model}


def _build_embedding_data_batch(texts: list, model: str, api_mode: str, dim: int) -> dict:
    """根据 api_mode 构建批量 embedding API 请求体"""
    if api_mode == "dashscope_multimodal":
        return {
            "model": model,
            "input": {"contents": [{"text": t} for t in texts]},
            "parameters": {"dimension": dim}
        }
    elif api_mode == "dashscope_text":
        return {
            "model": model,
            "input": {"texts": texts}
        }
    else:  # openai
        return {"input": texts, "model": model}


def _parse_embedding_response(result: dict, api_mode: str) -> list:
    """根据 api_mode 从响应中提取向量"""
    if api_mode == "dashscope_multimodal":
        return result["output"]["embeddings"][0]["embedding"]
    elif api_mode == "dashscope_text":
        return result["output"]["embeddings"][0]["embedding"]
    else:  # openai
        return result["data"][0]["embedding"]


def _parse_embedding_response_batch(result: dict, api_mode: str) -> list:
    """根据 api_mode 从批量响应中提取向量列表"""
    if api_mode in ("dashscope_multimodal", "dashscope_text"):
        return [item["embedding"] for item in result["output"]["embeddings"]]
    else:  # openai
        return [item["embedding"] for item in result["data"]]


# ===== 文档元数据管理 =====

META_FILE = DATA_DIR / "knowledge_meta.json"


def _load_meta() -> list:
    if META_FILE.exists():
        try:
            with open(META_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except (json.JSONDecodeError, IOError):
            return []
    return []


def _save_meta(meta: list):
    ensure_dirs()
    with open(META_FILE, "w", encoding="utf-8") as f:
        json.dump(meta, f, ensure_ascii=False, indent=2)


def ensure_dirs():
    KNOWLEDGE_DIR.mkdir(parents=True, exist_ok=True)
    (DATA_DIR / "chunks").mkdir(parents=True, exist_ok=True)


# ===== 文档列表 =====

def list_documents() -> list:
    meta = _load_meta()
    valid = []
    for doc in meta:
        doc_path = KNOWLEDGE_DIR / doc["filename"]
        if doc_path.exists():
            doc["size_display"] = _format_size(doc.get("size", 0))
            valid.append(doc)
    return valid


def _format_size(size_bytes: int) -> str:
    if size_bytes < 1024:
        return f"{size_bytes}B"
    elif size_bytes < 1024 * 1024:
        return f"{size_bytes/1024:.1f}KB"
    else:
        return f"{size_bytes/1024/1024:.1f}MB"


def get_document(doc_id: str) -> Optional[dict]:
    meta = _load_meta()
    for doc in meta:
        if doc["id"] == doc_id:
            return doc
    return None


# ===== 文档添加 =====

def add_document(filename: str, content: str) -> dict:
    """添加文档到知识库（含 sqlite-vec 向量索引）"""
    ensure_dirs()
    _ensure_db()
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    file_id = hashlib.md5(f"{filename}{now}".encode()).hexdigest()[:12]
    size = len(content.encode("utf-8"))

    # 保存原始文件
    doc_path = KNOWLEDGE_DIR / filename
    with open(doc_path, "w", encoding="utf-8") as f:
        f.write(content)

    # 切分文本
    chunks = _chunk_text(content)

    # 批量生成向量（API 调用）
    vectors = _get_embeddings_batch(chunks) if chunks else None

    # 存入 sqlite-vec
    conn = _get_vec_conn()
    for i, chunk in enumerate(chunks):
        chunk_id = f"{file_id}_{i}"

        # chunks 元数据
        conn.execute(
            "INSERT OR REPLACE INTO knowledge_chunks (id, doc_id, chunk_index, content, filename, title, created_at) VALUES (?, ?, ?, ?, ?, ?, ?)",
            (chunk_id, file_id, i, chunk, filename, Path(filename).stem, now)
        )

        # 向量（仅在 API 成功时）
        if vectors and i < len(vectors) and vectors[i] is not None:
            try:
                conn.execute(
                    "INSERT OR REPLACE INTO knowledge_vectors (id, vector) VALUES (?, ?)",
                    (chunk_id, _vec_to_str(vectors[i]))
                )
            except Exception as e:
                print(f"[sqlite-vec] 向量存储失败: {e}")

    conn.commit()
    conn.close()

    # 保存 chunks fallback（关键词搜索用）
    _save_chunks_fallback(file_id, chunks)

    # 保存元数据
    doc_info = {
        "id": file_id,
        "filename": filename,
        "title": Path(filename).stem,
        "size": size,
        "chunks": len(chunks),
        "created_at": now,
        "updated_at": now,
        "index_type": "vector" if _embedding_available else "keyword"
    }

    meta = _load_meta()
    for i, d in enumerate(meta):
        if d["filename"] == filename:
            meta[i] = doc_info
            break
    else:
        meta.append(doc_info)
    _save_meta(meta)

    return doc_info


# ===== 文本分块 =====

def _chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> list:
    """将长文本切分成块，带重叠"""
    text = text.strip()
    if not text:
        return []
    chunks = []
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    current_chunk = ""
    for para in paragraphs:
        if len(para) > chunk_size:
            sentences = [s.strip() for s in para.replace("。", "。\n").replace("！", "！\n")
                        .replace("？", "？\n").replace("；", "；\n").split("\n") if s.strip()]
            for sent in sentences:
                if len(current_chunk) + len(sent) > chunk_size:
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                    current_chunk = sent
                else:
                    current_chunk += sent + ""
        else:
            if len(current_chunk) + len(para) > chunk_size:
                chunks.append(current_chunk.strip())
                current_chunk = para
            else:
                current_chunk += para + "\n"
    if current_chunk.strip():
        chunks.append(current_chunk.strip())
    if not chunks:
        chunks = [text[:chunk_size]]
    return chunks


# ===== 降级：Chunks 文件存储（关键词搜索用） =====

CHUNKS_DIR = DATA_DIR / "chunks"


def _save_chunks_fallback(doc_id: str, chunks: list):
    CHUNKS_DIR.mkdir(parents=True, exist_ok=True)
    chunk_file = CHUNKS_DIR / f"{doc_id}.json"
    with open(chunk_file, "w", encoding="utf-8") as f:
        json.dump({"doc_id": doc_id, "chunks": chunks}, f, ensure_ascii=False)


def _load_chunks_fallback(doc_id: str) -> list:
    chunk_file = CHUNKS_DIR / f"{doc_id}.json"
    if chunk_file.exists():
        try:
            with open(chunk_file, "r", encoding="utf-8") as f:
                data = json.load(f)
                return data.get("chunks", [])
        except (json.JSONDecodeError, IOError):
            return []
    return []


# ===== 文档删除 =====

def delete_document(doc_id: str) -> bool:
    """删除文档（含向量索引）"""
    meta = _load_meta()
    doc = None
    for d in meta:
        if d["id"] == doc_id:
            doc = d
            break
    if not doc:
        return False

    # 删除原始文件
    doc_path = KNOWLEDGE_DIR / doc["filename"]
    if doc_path.exists():
        doc_path.unlink()

    # 删除 sqlite-vec 中的向量和 chunks 元数据
    try:
        conn = _get_vec_conn()
        conn.execute("DELETE FROM knowledge_chunks WHERE doc_id=?", (doc_id,))
        conn.execute("DELETE FROM knowledge_vectors WHERE id LIKE ?", (f"{doc_id}_%",))
        conn.commit()
        conn.close()
    except Exception as e:
        print(f"[sqlite-vec] 删除失败: {e}")

    # 删除 chunks fallback
    chunk_file = CHUNKS_DIR / f"{doc_id}.json"
    if chunk_file.exists():
        chunk_file.unlink()

    # 更新元数据
    meta = [d for d in meta if d["id"] != doc_id]
    _save_meta(meta)
    return True


# ===== 语义搜索（sqlite-vec） =====

def _semantic_search(query: str, top_k: int = 50) -> list:
    """使用 sqlite-vec 进行语义搜索"""
    global _embedding_available, _embedding_failure_time
    # 超时重试：如果上次失败超过 5 分钟，自动复位重试 API
    if not _embedding_available:
        if time.time() - _embedding_failure_time > _EMBEDDING_RETRY_INTERVAL:
            _embedding_available = True
        else:
            return [], "embedding_unavailable"
    _ensure_db()
    # _ensure_db 可能通过 _get_vec_conn 将 _embedding_available 设为 False
    if not _embedding_available:
        return [], "embedding_unavailable"

    query_vec = _get_embedding(query)
    if query_vec is None:
        return [], "embedding_failed"

    try:
        conn = _get_vec_conn()
        vec_str = _vec_to_str(query_vec)

        rows = conn.execute(
            f"""
            SELECT v.id, v.distance
            FROM knowledge_vectors v
            WHERE v.vector MATCH ?
              AND k = ?
            ORDER BY v.distance
            """,
            (vec_str, top_k)
        ).fetchall()

        conn.close()

        # 从主库获取 chunk 元数据（跨数据库查询）
        from database import get_db
        main_conn = get_db()
        results = []
        for rank, row in enumerate(rows):
            distance = row["distance"]
            score = max(0, 1.0 - distance / 2.0)
            chunk_id = row["id"]

            chunk_row = main_conn.execute(
                "SELECT content, doc_id, chunk_index, filename, title FROM knowledge_chunks WHERE id = ?",
                (chunk_id,)
            ).fetchone()

            results.append({
                "key": chunk_id,
                "doc_id": chunk_row["doc_id"] if chunk_row else "",
                "filename": chunk_row["filename"] if chunk_row else "",
                "title": chunk_row["title"] if chunk_row else "",
                "chunk_index": chunk_row["chunk_index"] if chunk_row else 0,
                "content": (chunk_row["content"] or "")[:300] if chunk_row else "",
                "semantic_rank": rank + 1,
                "semantic_score": round(score, 4)
            })

        main_conn.close()
        return results, "ok"

    except Exception as e:
        print(f"[sqlite-vec] 语义搜索失败: {e}")
        return [], f"error: {e}"


# ===== 混合检索（Hybrid Search） =====

def search_knowledge(query: str, top_k: int = 20) -> dict:
    """
    Hybrid Search（混合检索）：语义搜索 + 关键词搜索，RRF 加权融合
    
    当 embedding API 不可用时，自动降级为纯关键词搜索
    返回结果中会标注当前使用的搜索模式
    """
    if not query.strip():
        return {"results": [], "mode": "empty"}

    # 1. 语义搜索（sqlite-vec）
    semantic_results, semantic_status = _semantic_search(query, top_k=50)

    # 2. 关键词搜索（始终执行，作为降级保障）
    keyword_results = []
    try:
        meta = _load_meta()
        if meta:
            query_words = _tokenize(query)
            kw_all = []
            for doc in meta:
                chunks = _load_chunks_fallback(doc["id"])
                for idx, chunk in enumerate(chunks):
                    score = _score_chunk(chunk, query_words)
                    if score > 0:
                        kw_all.append({
                            "key": f"{doc['id']}_{idx}",
                            "doc_id": doc["id"],
                            "filename": doc["filename"],
                            "title": doc["title"],
                            "chunk_index": idx,
                            "content": chunk[:300],
                            "keyword_score": score
                        })
            kw_all.sort(key=lambda x: x["keyword_score"], reverse=True)
            for rank, item in enumerate(kw_all[:50]):
                item["keyword_rank"] = rank + 1
                keyword_results.append(item)
    except Exception as e:
        print(f"[关键词搜索失败]: {e}")

    # 确定搜索模式
    if semantic_status != "ok" or not semantic_results:
        # 降级为纯关键词搜索
        for item in keyword_results[:top_k]:
            item["score"] = round(item.get("keyword_score", 0), 4)
        reason = ""
        if semantic_status == "embedding_unavailable":
            reason = "向量索引不可用（知识库尚未建立向量索引）"
        elif semantic_status == "embedding_failed":
            reason = "向量服务不可用（API调用失败）"
        elif not keyword_results:
            reason = "向量+关键词均未匹配到结果"
        elif semantic_status != "ok":
            reason = f"向量搜索失败（{semantic_status}），已降级为关键词搜索"
        return {
            "results": keyword_results[:top_k],
            "mode": "keyword_only",
            "embedding_available": _embedding_available,
            "degraded_reason": reason
        }

    # 3. RRF 融合（仅当语义搜索成功时）
    K = 60
    rrf_scores = {}

    for item in semantic_results:
        key = item["key"]
        rrf_scores[key] = {
            "doc_id": item["doc_id"],
            "filename": item["filename"],
            "title": item["title"],
            "chunk_index": item["chunk_index"],
            "content": item["content"],
            "score": 1.0 / (K + item["semantic_rank"])
        }

    for item in keyword_results:
        key = item["key"]
        if key in rrf_scores:
            rrf_scores[key]["score"] += 1.0 / (K + item["keyword_rank"])
        else:
            rrf_scores[key] = {
                "doc_id": item["doc_id"],
                "filename": item["filename"],
                "title": item["title"],
                "chunk_index": item["chunk_index"],
                "content": item["content"],
                "score": 1.0 / (K + item["keyword_rank"])
            }

    # 4. 按 RRF 分数排序
    final = sorted(rrf_scores.values(), key=lambda x: x["score"], reverse=True)
    
    # 5. 精排优化（Rerank）：对候选结果做关键词匹配度 + 位置 + 密度 综合加权
    try:
        query_lower = query.lower()
        # 提取查询关键词
        stop_words = {'的', '了', '在', '是', '我', '有', '和', '就', '不', '人', '都', '一', '一个', '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看', '好', '自己', '这', '他', '她', '它', '们', '那', '些', '什么', '怎么', '如何', '为什么', '可以', '吗', '吧', '呢', '啊', '哦', '嗯'}
        query_keywords = [w for w in query_lower.split() if w not in stop_words and len(w) > 1]
        if query_keywords:
            for item in final[:top_k]:
                content = item.get("content", "")
                content_lower = content.lower()
                # (1) 精确匹配度：完整查询短语出现的次数
                exact_count = content_lower.count(query_lower)
                # (2) 关键词匹配度：命中多少个关键词
                match_count = sum(1 for kw in query_keywords if kw in content_lower)
                # (3) 关键词密度：关键词字符占总内容的比率
                matched_chars = sum(len(kw) for kw in query_keywords if kw in content_lower)
                density = matched_chars / max(len(content), 1)
                # (4) 位置加分：关键词越靠前越好
                first_pos = len(content)
                for kw in query_keywords:
                    pos = content_lower.find(kw)
                    if pos >= 0 and pos < first_pos:
                        first_pos = pos
                position_bonus = max(0, 1.0 - first_pos / max(len(content), 1))
                
                # 综合精排分数
                rerank_score = 0
                rerank_score += exact_count * 0.3                    # 精确匹配加分
                rerank_score += (match_count / len(query_keywords)) * 0.4  # 关键词覆盖率
                rerank_score += density * 0.2                         # 密度加分
                rerank_score += position_bonus * 0.1                  # 位置加分
                
                if rerank_score > 0:
                    item["score"] = round(item["score"] + rerank_score, 4)
                    item["rerank_score"] = round(rerank_score, 4)
            # 重新排序
            final = sorted(final, key=lambda x: x["score"], reverse=True)
    except Exception:
        pass
    
    for item in final:
        item["score"] = round(item["score"], 4)

    return {
        "results": final[:top_k],
        "mode": "hybrid",
        "embedding_available": True
    }


# ===== 纯关键词搜索（降级用） =====

def _tokenize(text: str) -> list:
    """中文分词（二元组 + 整词）"""
    import re
    words = re.findall(r'[\u4e00-\u9fa5a-zA-Z0-9]+', text.lower())
    all_tokens = []
    for w in words:
        if re.match(r'^[\u4e00-\u9fa5]+$', w):
            for i in range(len(w) - 1):
                all_tokens.append(w[i:i+2])
            all_tokens.append(w)
        else:
            all_tokens.append(w)
    return list(set(all_tokens))


def _score_chunk(chunk: str, query_words: list) -> float:
    """关键词匹配打分"""
    if not query_words:
        return 0
    chunk_lower = chunk.lower()
    score = 0
    for word in query_words:
        count = chunk_lower.count(word)
        if count > 0:
            score += count * (1 + 1 / (len(chunk) + 1))
    return score


# ===== 客户图片管理 =====

CUSTOMER_IMAGES_DIR = DATA_DIR / "customer_images"


def ensure_image_dirs():
    """确保图片目录存在"""
    CUSTOMER_IMAGES_DIR.mkdir(parents=True, exist_ok=True)


def init_image_table():
    """初始化图片元数据表"""
    _ensure_db()
    conn = _get_vec_conn()
    conn.execute("""
        CREATE TABLE IF NOT EXISTS customer_images (
            id TEXT PRIMARY KEY,
            customer_id INTEGER NOT NULL,
            filename TEXT NOT NULL,
            original_name TEXT DEFAULT '',
            description TEXT DEFAULT '',
            filepath TEXT NOT NULL,
            created_at TEXT DEFAULT (datetime('now','localtime'))
        )
    """)
    try:
        conn.execute(f"""
            CREATE VIRTUAL TABLE IF NOT EXISTS image_vectors USING vec0(
                id TEXT PRIMARY KEY,
                vector FLOAT[{_embedding_dim}]
            )
        """)
    except Exception:
        pass
    conn.commit()
    conn.close()


def add_customer_image(customer_id: int, filename: str, content_bytes: bytes, description: str = "") -> dict:
    """保存客户图片并生成描述和向量索引"""
    import hashlib
    from datetime import datetime
    
    ensure_image_dirs()
    init_image_table()
    
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    img_id = hashlib.md5(f"{filename}{now}".encode()).hexdigest()[:12]
    
    # 保存图片文件
    ext = Path(filename).suffix.lower()
    save_name = f"c{customer_id}_{img_id}{ext}"
    save_path = CUSTOMER_IMAGES_DIR / save_name
    save_path.write_bytes(content_bytes)
    
    # 如果没有描述，调用视觉模型生成
    if not description:
        description = _describe_image(str(save_path))
    
    # 生成描述向量
    vector = _get_embedding(description) if description else None
    
    # 存入数据库
    conn = _get_vec_conn()
    conn.execute(
        "INSERT OR REPLACE INTO customer_images (id, customer_id, filename, original_name, description, filepath, created_at) VALUES (?, ?, ?, ?, ?, ?, ?)",
        (img_id, customer_id, save_name, filename, description, str(save_path), now)
    )
    
    # 存向量
    if vector:
        try:
            conn.execute(
                "INSERT OR REPLACE INTO image_vectors (id, vector) VALUES (?, ?)",
                (img_id, _vec_to_str(vector))
            )
        except Exception as e:
            print(f"[图片] 向量存储失败: {e}")
    
    conn.commit()
    conn.close()
    
    return {
        "id": img_id,
        "customer_id": customer_id,
        "filename": save_name,
        "original_name": filename,
        "description": description,
        "filepath": str(save_path)
    }


def _describe_image(image_path: str) -> str:
    """调用视觉模型描述图片内容 - 多provider轮询，支持过期自动跳过"""
    import base64
    from pathlib import Path
    import json
    import urllib.request
    import urllib.error
    from datetime import datetime

    with open(image_path, "rb") as f:
        b64 = base64.b64encode(f.read()).decode("utf-8")
    ext = Path(image_path).suffix.lower()
    mime = "image/jpeg" if ext in (".jpg", ".jpeg") else "image/png"
    # 从配置读取图片识别提示词模板
    try:
        import config_manager as _cfg
        _prompts = _cfg.get_prompts()
        prompt = _prompts.get("image_recognition_chat", "")
        if not prompt:
            raise ValueError("empty prompt")
    except Exception:
        prompt = """只输出6行，每行一个指标值，不要任何分析、解释、额外文字：

    【油脂量】少量/中等/大量/满层
    【油脂状态】绵密/一片一片/油花/油滴/粘稠/厚油层/油水混合
    【油脂颜色】黄色/棕色/红色/黑色/橙色/透明/褐色
    【大便量】少量/中等/大量/无
    【大便状态】稀水状/不成形糊状/软便/成形条状/干硬/粘稠/无
    【大便颜色】黄色/棕色/绿色/黑色/红色/褐色/无

    严格按以上格式输出6行，不要分析、不要思考过程、不要额外说明。如果图片不是排便图，根据实际内容用类似格式描述。"""

    def _is_expired(expires_at: str) -> bool:
        """检查是否已过期"""
        if not expires_at:
            return False
        try:
            expire = datetime.fromisoformat(expires_at)
            return datetime.now() >= expire
        except Exception:
            return False

    def _try_call(model: str, api_key: str, base_url: str, is_reasoning: bool = False) -> str:
        """调用单个视觉模型，返回描述文本或抛出异常"""
        import time
        t0 = time.time()
        try:
            data = json.dumps({
                "model": model,
                "messages": [{
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}}
                    ]
                }],
                "max_tokens": 500
            }).encode("utf-8")
            req = urllib.request.Request(
                f"{base_url.rstrip('/')}/chat/completions",
                data=data,
                headers={
                    "Authorization": f"Bearer {api_key}",
                    "Content-Type": "application/json"
                },
                method="POST"
            )
            resp = urllib.request.urlopen(req, timeout=60)
            result = json.loads(resp.read().decode("utf-8"))
            # 优先取 content，没有则取 reasoning
            try:
                text = result["choices"][0]["message"]["content"].strip()
            except (KeyError, IndexError):
                text = result["choices"][0]["message"].get("reasoning", "").strip()
            if not text:
                text = result["choices"][0]["message"].get("reasoning_content", "").strip()
            elapsed = time.time() - t0
            # 写日志
            try:
                import logger
                logger.log_event("vision", "describe", model=model, success=True, duration=round(elapsed, 2))
            except Exception:
                pass
            return text
        except Exception as e:
            elapsed = time.time() - t0
            try:
                import logger
                logger.log_event("vision", "describe", model=model, success=False, duration=round(elapsed, 2), error=str(e)[:100])
            except Exception:
                pass
            raise

    errors = []
    config = None

    # === 尝试1: 从配置读取视觉模型（优先使用 vision_default_id）===
    try:
        import config_manager as _cfg
        v_cfg = _cfg.get_vision_model_config()
        if v_cfg and v_cfg.get("api_key"):
            desc = _try_call(v_cfg["model"], v_cfg["api_key"], v_cfg["base_url"].rstrip("/"))
            if desc:
                print(f"[图片描述] {v_cfg['model']} 成功: {desc[:50]}...")
                return desc
            else:
                errors.append(f"{v_cfg['model']}: 返回空")
        else:
            errors.append("get_vision_model_config(): 无可用配置")
    except urllib.error.HTTPError as e:
        err_body = e.read().decode("utf-8", errors="replace")[:100]
        errors.append(f"vision_default: HTTP {e.code} {err_body}")
    except Exception as e:
        errors.append(f"vision_default: {e}")

    # === 尝试2: 遍历 models.list 中的其他视觉模型 ===
        try:
            import config_manager as _cfg
            _all_cfg = _cfg.load_config()
            _models = _all_cfg.get("models", {}).get("list", {})
            _def_id = _all_cfg.get("models", {}).get("vision_default_id", "")
            for _mid, _mcfg in _models.items():
                if _mid == _def_id:
                    continue
                _mname = _mcfg.get("model", "")
                _cats = _mcfg.get("categories", [])
                # 判断是否视觉模型：检查 categories 或模型名
                is_vision_model = "vision" in _cats or \
                                  ("-4v-" in _mname.lower() or "-1v-" in _mname.lower() or "vision" in _mname.lower()) or \
                                  (_mname.lower().startswith("step-") and "flash" in _mname.lower())
                if is_vision_model:
                    _key = _mcfg.get("api_key", "")
                    _url = _mcfg.get("base_url", "").rstrip("/")
                    if _key and _url:
                        try:
                            desc = _try_call(_mname, _key, _url)
                            if desc:
                                print(f"[图片描述] {_mname} 成功: {desc[:50]}...")
                                return desc
                        except Exception as _e:
                            errors.append(f"{_mname}: {_e}")
        except Exception as e:
            errors.append(f"models.list遍历: {e}")

        # === 尝试3: 使用默认话术模型（如果它有视觉能力）===
        try:
            import config_manager as _cfg
            _llm_cfg = _cfg.get_llm_config()
            if _llm_cfg and _llm_cfg.get("api_key") and _llm_cfg.get("model"):
                _mname = _llm_cfg.get("model", "")
                # 跳过上面已经试过的模型
                if _mname not in [e.split(":")[0] for e in errors]:
                    try:
                        desc = _try_call(_mname, _llm_cfg["api_key"], _llm_cfg["base_url"].rstrip("/"))
                        if desc:
                            print(f"[图片描述] {_mname}（默认话术模型）成功: {desc[:50]}...")
                            return desc
                    except Exception as _e:
                        errors.append(f"{_mname}: {_e}")
        except Exception as e:
            errors.append(f"默认话术模型: {e}")

        # === 全部失败 ===
    err_msg = "；".join(errors) if errors else "所有视觉模型均不可用"
    print(f"[图片描述] 全部失败: {err_msg}")
    return f"[识别失败] {err_msg}"


def search_customer_images(customer_id: int, query: str = "") -> list:
    """搜索客户相关的图片（语义搜索）"""
    init_image_table()
    conn = _get_vec_conn()
    
    if not query.strip():
        # 无查询词，返回该客户最新图片
        rows = conn.execute(
            "SELECT * FROM customer_images WHERE customer_id=? ORDER BY created_at DESC LIMIT 10",
            (customer_id,)
        ).fetchall()
        conn.close()
        return [dict(r) for r in rows]
    
    # 语义搜索
    query_vec = _get_embedding(query)
    if query_vec:
        try:
            vec_str = _vec_to_str(query_vec)
            rows = conn.execute(
                f"""
                SELECT v.id, v.distance, c.customer_id, c.filename, c.original_name, c.description, c.filepath, c.created_at
                FROM image_vectors v
                LEFT JOIN customer_images c ON v.id = c.id
                WHERE c.customer_id = ?
                  AND v.vector MATCH ?
                  AND k = 10
                ORDER BY v.distance
                """,
                (customer_id, vec_str)
            ).fetchall()
            conn.close()
            results = []
            for r in rows:
                item = dict(r)
                item["similarity"] = round(max(0, 1.0 - item.get("distance", 1) / 2), 3)
                results.append(item)
            return results
        except Exception as e:
            print(f"[图片搜索] 向量搜索失败: {e}")
    
    # 降级：返回最新图片
    conn.close()
    rows = conn.execute(
        "SELECT * FROM customer_images WHERE customer_id=? ORDER BY created_at DESC LIMIT 5",
        (customer_id,)
    ).fetchall()
    return [dict(r) for r in rows]


def get_customer_images(customer_id: int) -> list:
    """获取客户的所有图片"""
    init_image_table()
    conn = _get_vec_conn()
    rows = conn.execute(
        "SELECT * FROM customer_images WHERE customer_id=? ORDER BY created_at DESC",
        (customer_id,)
    ).fetchall()
    conn.close()
    return [dict(r) for r in rows]


def delete_customer_image(img_id: str) -> bool:
    """删除客户图片"""
    init_image_table()
    conn = _get_vec_conn()
    row = conn.execute("SELECT filepath FROM customer_images WHERE id=?", (img_id,)).fetchone()
    if not row:
        conn.close()
        return False
    
    # 删除文件
    try:
        os.unlink(row["filepath"])
    except Exception:
        pass
    
    conn.execute("DELETE FROM customer_images WHERE id=?", (img_id,))
    try:
        conn.execute("DELETE FROM image_vectors WHERE id=?", (img_id,))
    except Exception:
        pass
    conn.commit()
    conn.close()
    return True

def get_knowledge_status() -> dict:
    """获取知识库状态"""
    _ensure_db()
    meta = _load_meta()
    total_chunks = sum(d.get("chunks", 0) for d in meta)
    total_size = sum(d.get("size", 0) for d in meta)

    # 检查 sqlite-vec 向量数量
    vec_count = 0
    try:
        conn = _get_vec_conn()
        row = conn.execute("SELECT COUNT(*) as cnt FROM knowledge_vectors").fetchone()
        vec_count = row["cnt"] if row else 0
        conn.close()
    except Exception:
        vec_count = total_chunks

    return {
        "total_documents": len(meta),
        "total_chunks": vec_count if vec_count > 0 else total_chunks,
        "total_size": total_size,
        "total_size_display": _format_size(total_size),
        "index_type": "vector" if _embedding_available else "keyword",
        "embedding_available": _embedding_available,
        "vec_count": vec_count
    }


def get_document_content(doc_id: str) -> Optional[str]:
    meta = _load_meta()
    for d in meta:
        if d["id"] == doc_id:
            doc_path = KNOWLEDGE_DIR / d["filename"]
            if doc_path.exists():
                try:
                    return doc_path.read_text(encoding="utf-8")
                except Exception:
                    return None
    return None


# ===== 重建索引 =====

def reindex_all():
    """重新索引所有文档（重建向量）"""
    meta = _load_meta()
    if not meta:
        return True

    # 重建向量索引表（DROP旧表，用当前维度重建）
    try:
        conn = _get_vec_conn()
        conn.execute("DROP TABLE IF EXISTS knowledge_vectors")
        conn.execute(f"""
            CREATE VIRTUAL TABLE IF NOT EXISTS knowledge_vectors USING vec0(
                id TEXT PRIMARY KEY,
                vector FLOAT[{_embedding_dim}]
            )
        """)
        conn.commit()
        conn.close()
        print(f"[重建索引] 向量表已重建 (FLOAT[{_embedding_dim}])")
    except Exception as e:
        print(f"[重建索引] 重建向量表失败: {e}")
        return False

    # 重新索引每个文档
    for doc in meta:
        doc_path = KNOWLEDGE_DIR / doc["filename"]
        if doc_path.exists():
            try:
                content = doc_path.read_text(encoding="utf-8")
                chunks = _chunk_text(content)
                doc["chunks"] = len(chunks)
                doc["size"] = len(content.encode("utf-8"))

                # 批量生成向量
                vectors = _get_embeddings_batch(chunks) if chunks else None

                # 写入主库 chunks 表
                mconn = sqlite3.connect(str(DB_PATH))
                mconn.row_factory = sqlite3.Row
                mconn.execute("DELETE FROM knowledge_chunks WHERE doc_id=?", (doc["id"],))
                now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                title = doc.get("title", "") or ""
                for i, chunk in enumerate(chunks):
                    chunk_id = f"{doc['id']}_{i}"
                    mconn.execute(
                        "INSERT OR REPLACE INTO knowledge_chunks (id, doc_id, chunk_index, content, filename, title, created_at) VALUES (?, ?, ?, ?, ?, ?, ?)",
                        (chunk_id, doc["id"], i, chunk, doc["filename"], title, now)
                    )
                mconn.commit()
                mconn.close()

                # 写入向量库
                if vectors:
                    vconn = _get_vec_conn()
                    for i, chunk in enumerate(chunks):
                        if i < len(vectors) and vectors[i] is not None:
                            try:
                                vconn.execute(
                                    "INSERT OR REPLACE INTO knowledge_vectors (id, vector) VALUES (?, ?)",
                                    (f"{doc['id']}_{i}", _vec_to_str(vectors[i]))
                                )
                            except Exception:
                                pass
                    vconn.commit()
                    vconn.close()

                doc["index_type"] = "vector" if _embedding_available else "keyword"

            except Exception as e:
                print(f"[重建索引] {doc['filename']} 失败: {e}")

    _save_meta(meta)
    return True


# ===== 文件上传处理 =====

def save_uploaded_file(filename: str, content: str) -> str:
    ensure_dirs()
    filepath = KNOWLEDGE_DIR / filename
    if isinstance(content, str):
        filepath.write_text(content, encoding="utf-8")
    else:
        filepath.write_bytes(content)
    return str(filepath)


def read_docx(filepath: str) -> str:
    try:
        from docx import Document
        doc = Document(filepath)
        text = []
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    text.append(" | ".join(cells))
        return "\n".join(text)
    except Exception as e:
        return f"[读取文档失败: {e}]"


def read_pdf(filepath: str) -> str:
    try:
        import PyPDF2
        with open(filepath, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            text = []
            for page in reader.pages:
                t = page.extract_text()
                if t.strip():
                    text.append(t.strip())
        return "\n".join(text)
    except ImportError:
        return "[PDF阅读器未安装]"
    except Exception as e:
        return f"[读取PDF失败: {e}]"


def extract_text_from_file(filepath: str) -> str:
    ext = Path(filepath).suffix.lower()
    try:
        if ext == ".txt":
            with open(filepath, "r", encoding="utf-8", errors="replace") as f:
                return f.read()
        elif ext == ".docx":
            return read_docx(filepath)
        elif ext == ".pdf":
            return read_pdf(filepath)
        elif ext in (".md", ".csv", ".json", ".xml", ".yaml", ".yml"):
            with open(filepath, "r", encoding="utf-8", errors="replace") as f:
                return f.read()
        else:
            return f"[不支持的文档格式: {ext}]"
    except Exception as e:
        return f"[读取文件失败: {e}]"